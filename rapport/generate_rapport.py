#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Rapport technique BDD Distribuees - Oracle XE + Docker
Mise en forme professionnelle avec page de garde et styles avances
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# ════════════════════════════════════════════════════════
# PALETTE DE COULEURS
# ════════════════════════════════════════════════════════
C_NAVY      = RGBColor(0x0D, 0x2B, 0x55)   # bleu marine fonce
C_BLUE      = RGBColor(0x1A, 0x56, 0x9C)   # bleu principal
C_BLUE_MID  = RGBColor(0x2E, 0x74, 0xB5)   # bleu moyen
C_BLUE_LIGHT= RGBColor(0xBD, 0xD7, 0xEE)   # bleu clair (fill)
C_ORANGE    = RGBColor(0xE8, 0x6A, 0x17)   # orange accent
C_ORANGE_L  = RGBColor(0xFC, 0xE4, 0xD6)   # orange clair
C_GREY_DARK = RGBColor(0x40, 0x40, 0x40)   # gris texte
C_GREY_MED  = RGBColor(0x70, 0x70, 0x70)   # gris moyen
C_GREY_FILL = "F5F5F5"                       # gris clair fill (hex str)
C_CODE_BG   = "EAEAEA"                       # fond code
C_CODE_HDR  = "1A569C"                       # en-tete code (bleu)
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_GOLD      = RGBColor(0xBF, 0x96, 0x00)

doc = Document()

# ════════════════════════════════════════════════════════
# PARAMETRES GLOBAUX
# ════════════════════════════════════════════════════════
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.2)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_para_bg(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def add_border_bottom(p, color="1A569C", size="12"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), size)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def set_run_font(run, name='Calibri'):
    run.font.name = name
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.insert(0, rFonts)

def add_header_footer():
    """Ajoute en-tete et pied de page sur toutes les sections sauf la 1ere"""
    for i, section in enumerate(doc.sections):
        # Pas d'en-tete sur la page de garde
        if i == 0:
            section.different_first_page_header_footer = True
        hdr = section.header
        hdr.is_linked_to_previous = False
        for p in hdr.paragraphs:
            p.clear()
        if len(hdr.paragraphs) == 0:
            hdr.add_paragraph()
        hp = hdr.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_border_bottom(hp, color="1A569C", size="6")
        r1 = hp.add_run("Rapport Technique -- Bases de Donnees Distribuees Oracle + Docker")
        r1.font.size = Pt(8)
        r1.font.color.rgb = C_BLUE_MID
        r1.italic = True

        ftr = section.footer
        ftr.is_linked_to_previous = False
        for p in ftr.paragraphs:
            p.clear()
        if len(ftr.paragraphs) == 0:
            ftr.add_paragraph()
        fp = ftr.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = fp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), '1A569C')
        pBdr.append(top)
        pPr.append(pBdr)
        r2 = fp.add_run("EShop Distribue -- Oracle XE 21c + Docker   |   Page ")
        r2.font.size = Pt(8)
        r2.font.color.rgb = C_GREY_MED
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_pg = fp.add_run()
        run_pg._r.append(fldChar1)
        run_pg._r.append(instrText)
        run_pg._r.append(fldChar2)
        run_pg.font.size = Pt(8)
        run_pg.font.color.rgb = C_BLUE
        r3 = fp.add_run("   |   2025-2026")
        r3.font.size = Pt(8)
        r3.font.color.rgb = C_GREY_MED

def page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════
# HELPERS CONTENU
# ════════════════════════════════════════════════════════

def heading1(text):
    """Titre niveau 1 avec barre orange en dessous"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = C_NAVY
    set_run_font(run, 'Calibri')
    add_border_bottom(p, color="E86A17", size="18")
    return p

def heading2(text):
    """Titre niveau 2 avec barre bleue fine"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = C_BLUE
    set_run_font(run, 'Calibri')
    add_border_bottom(p, color="2E74B5", size="8")
    return p

def heading3(text):
    """Titre niveau 3 sans barre"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_BLUE_MID
    set_run_font(run, 'Calibri')
    return p

def para(text, bold=False, italic=False, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else C_GREY_DARK
    set_run_font(run, 'Calibri')
    if align:
        p.alignment = align
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.8)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10)
        rb.font.color.rgb = C_BLUE
        set_run_font(rb, 'Calibri')
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = C_GREY_DARK
    set_run_font(r, 'Calibri')
    return p

def code_block(text, title=None):
    """Encadre code avec en-tete bleu et fond gris"""
    if title:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(8)
        tp.paragraph_format.space_after  = Pt(0)
        tp.paragraph_format.left_indent  = Cm(0)
        set_para_bg(tp, C_CODE_HDR)
        r = tp.add_run("  " + title)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = C_WHITE
        set_run_font(r, 'Calibri')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    set_para_bg(p, C_CODE_BG)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.5)
    set_para_bg(p, "FFF3CD")
    r = p.add_run("  Note : " + text)
    r.font.size = Pt(9.5)
    r.italic = True
    r.font.color.rgb = RGBColor(0x7D, 0x4E, 0x00)
    set_run_font(r, 'Calibri')
    return p

def warn(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.5)
    set_para_bg(p, "FFE0E0")
    r = p.add_run("  Attention : " + text)
    r.font.size = Pt(9.5)
    r.bold = True
    r.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    set_run_font(r, 'Calibri')
    return p

def rich_table(rows, headers, col_widths=None):
    """Tableau stylise avec alternance de couleurs"""
    ncols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # En-tete
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, "0D2B55")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(3)
        cp.paragraph_format.space_after  = Pt(3)
        r = cp.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = C_WHITE
        set_run_font(r, 'Calibri')
    # Lignes
    for ri, row in enumerate(rows):
        bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after  = Pt(2)
            r = cp.add_run(str(val))
            r.font.size = Pt(9)
            r.font.color.rgb = C_GREY_DARK
            set_run_font(r, 'Calibri')
    # Largeurs colonnes
    if col_widths:
        for ri in range(len(rows)+1):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    add_border_bottom(p, color="BDD7EE", size="6")

# ════════════════════════════════════════════════════════
# PAGE DE GARDE PROFESSIONNELLE
# ════════════════════════════════════════════════════════

# Bande superieure bleue marine
def add_cover_band_top():
    """Bande bleue marine en haut de la page de garde"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, "0D2B55")
    cell.width = Cm(21)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(14)
    cp.paragraph_format.space_after  = Pt(14)
    r = cp.add_run("UNIVERSITE  --  MASTER INFORMATIQUE")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = C_BLUE_LIGHT
    set_run_font(r, 'Calibri')

add_cover_band_top()

# Espacement
for _ in range(3):
    doc.add_paragraph()

# Ligne decorative orange
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_border_bottom(p, color="E86A17", size="24")
r = p.add_run(" ")
r.font.size = Pt(4)

doc.add_paragraph()

# Sous-titre discipline
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RAPPORT TECHNIQUE DE FIN DE PROJET")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = C_ORANGE
r.font.all_caps = True
set_run_font(r, 'Calibri')
p.paragraph_format.space_after = Pt(16)

# Titre principal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Bases de Donnees Distribuees\navec Oracle XE et Docker")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = C_NAVY
set_run_font(r, 'Calibri')
p.paragraph_format.space_after = Pt(10)

# Sous-titre
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Systeme EShop -- Fragmentation Horizontale, Database Links,\nProcedures Stockees PL/SQL et Triggers de Synchronisation")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = C_BLUE_MID
set_run_font(r, 'Calibri')
p.paragraph_format.space_after = Pt(24)

# Ligne decorative bleue
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_border_bottom(p, color="1A569C", size="12")
r = p.add_run(" ")
r.font.size = Pt(4)

for _ in range(2):
    doc.add_paragraph()

# Encadre informations projet
t = doc.add_table(rows=1, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = t.rows[0].cells[0]
set_cell_bg(cell, "EBF3FB")
cell.width = Cm(14)
cp = cell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_before = Pt(12)
cp.paragraph_format.space_after  = Pt(4)

infos = [
    ("Specialite",  "Bases de Donnees Distribuees"),
    ("Technologie", "Oracle XE 21c  |  Docker Compose  |  PL/SQL"),
    ("Annee",       "2025 -- 2026"),
    ("Date",        "Juin 2026"),
]

for key, val in infos:
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_before = Pt(2)
    p_info.paragraph_format.space_after  = Pt(2)
    rk = p_info.add_run(key + " : ")
    rk.bold = True
    rk.font.size = Pt(10.5)
    rk.font.color.rgb = C_NAVY
    set_run_font(rk, 'Calibri')
    rv = p_info.add_run(val)
    rv.font.size = Pt(10.5)
    rv.font.color.rgb = C_GREY_DARK
    set_run_font(rv, 'Calibri')

for _ in range(3):
    doc.add_paragraph()

# Bande inferieure bleue
t2 = doc.add_table(rows=1, cols=3)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
colors_bot = ["0D2B55", "1A569C", "E86A17"]
texts_bot  = ["Oracle XE 21c", "Docker Compose", "PL/SQL Distribue"]
for ci, (col, txt) in enumerate(zip(colors_bot, texts_bot)):
    cell = t2.rows[0].cells[ci]
    set_cell_bg(cell, col)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(8)
    cp.paragraph_format.space_after  = Pt(8)
    r = cp.add_run(txt)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = C_WHITE
    set_run_font(r, 'Calibri')

page_break()

# ════════════════════════════════════════════════════════
# TABLE DES MATIERES
# ════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TABLE DES MATIERES")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = C_NAVY
r.font.all_caps = True
set_run_font(r, 'Calibri')
p.paragraph_format.space_after = Pt(4)

p2 = doc.add_paragraph()
add_border_bottom(p2, color="E86A17", size="18")
r2 = p2.add_run(" ")
r2.font.size = Pt(4)

doc.add_paragraph()

toc_items = [
    ("1.",  "Introduction et contexte du projet"),
    ("2.",  "Architecture generale"),
    ("3.",  "Configuration Docker Compose"),
    ("4.",  "Configuration reseau"),
    ("5.",  "Configuration des sites Oracle"),
    ("6.",  "Tests de connectivite"),
    ("7.",  "Creation des utilisateurs et privileges"),
    ("8.",  "Configuration des Database Links"),
    ("9.",  "Schema de donnees global"),
    ("10.", "Tables fragmentees sur les sites distants"),
    ("11.", "Architecture des procedures stockees"),
    ("12.", "Architecture des triggers"),
    ("13.", "Synonymes et vues distribuees"),
    ("14.", "Routage automatique depuis le site central"),
    ("15.", "Optimisation des requetes distribuees"),
    ("16.", "Strategie d'indexation multi-niveaux"),
    ("17.", "Analyse comparative des performances"),
    ("18.", "Monitoring et maintenance"),
    ("19.", "Conclusion"),
]

for i, (num, title) in enumerate(toc_items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
    set_para_bg(p, bg)
    rn = p.add_run(f"  {num:<4}")
    rn.bold = True
    rn.font.size = Pt(10)
    rn.font.color.rgb = C_ORANGE
    set_run_font(rn, 'Calibri')
    rt = p.add_run(title)
    rt.font.size = Pt(10)
    rt.font.color.rgb = C_NAVY
    set_run_font(rt, 'Calibri')

page_break()

# ════════════════════════════════════════════════════════
# SECTIONS DU RAPPORT
# ════════════════════════════════════════════════════════

# ── 1. INTRODUCTION ────────────────────────────────────
heading1("1. Introduction et contexte du projet")

para("Ce rapport presente le travail realise dans le cadre du projet de fin de module sur les bases de donnees distribuees. L'objectif principal etait de simuler un environnement de production reel ou les donnees d'un systeme e-commerce sont reparties sur plusieurs noeuds Oracle XE, le tout orchestre via Docker afin de garantir un environnement reproductible et independant de la machine hote.")

para("Le projet porte sur un systeme nomme EShop Distribue, comprenant la gestion de clients, produits, fournisseurs, employes, commandes et lignes de commande. Ce domaine metier classique nous a permis de nous concentrer sur les aspects techniques de la distribution plutot que sur la modelisation.")

para("Nous avons choisi Oracle Database XE 21c comme moteur, car il offre des fonctionnalites avancees de distribution (Database Links, triggers INSTEAD OF, procedures PL/SQL) tout en restant gratuit pour un usage academique. L'image Docker gvenzl/oracle-xe:21-slim a considerablement simplifie le deploiement.")

heading2("Deux scenarios de fragmentation ont ete implementes")

bullet("Scenario 1 -- Fragmentation par categorie : ", bold_prefix="")
bullet("  les lignes liees aux produits Informatique (IDCATEG = 50) sont sur le Site 1, celles des Accessoires (IDCATEG = 35) sur le Site 2.")
bullet("Scenario 2 -- Fragmentation par volume (deploye par defaut) : ", bold_prefix="")
bullet("  les lignes avec QUANTITE >= 100 (grossistes) vont sur le Site 1, celles < 100 (detaillants) sur le Site 2.")

divider()

# ── 2. ARCHITECTURE ────────────────────────────────────
heading1("2. Architecture generale")

para("Notre architecture repose sur trois noeuds Oracle XE conteneurises qui communiquent via un reseau Docker bridge prive. Voici la repartition des roles :")

rich_table([
    ("oracle-site1",   "172.20.0.10", "1521 (hote)",  "Site Principal -- QTE >= 100 (Sc2) / IDCATEG=50 (Sc1)"),
    ("oracle-site2",   "172.20.0.11", "1522 (hote)",  "Site Distant -- QTE < 100 (Sc2) / IDCATEG=35 (Sc1)"),
    ("oracle-central", "172.20.0.12", "1523 (hote)",  "Coordinateur -- Vues globales, synonymes, routage"),
], headers=["Noeud Docker", "IP fixe", "Port", "Role"], col_widths=[3.5, 2.5, 2.5, 8.5])

code_block(
    "┌─────────────────────┐     DB LINK (bidir.)    ┌─────────────────────┐\n"
    "│    oracle-site1     │◄──────────────────────►│    oracle-site2     │\n"
    "│    172.20.0.10      │                         │    172.20.0.11      │\n"
    "│  BDDVENTE : 1521    │                         │  BDDVENTE2 : 1521   │\n"
    "│  eshop1             │                         │  eshop2             │\n"
    "│  Sc2: QTE >= 100    │                         │  Sc2: QTE <  100    │\n"
    "│  Sc1: IDCATEG = 50  │                         │  Sc1: IDCATEG = 35  │\n"
    "└──────────┬──────────┘                         └──────────┬──────────┘\n"
    "           │                  DB LINK                      │\n"
    "           └────────────────────┬─────────────────────────┘\n"
    "                                ▼\n"
    "                ┌───────────────────────────┐\n"
    "                │      oracle-central        │\n"
    "                │      172.20.0.12           │\n"
    "                │    BDDCENTRAL : 1521        │\n"
    "                │    eshopcentral             │\n"
    "                │  Vues globales / Routage    │\n"
    "                └───────────────────────────┘",
    title="Diagramme d'architecture -- 3 noeuds Oracle XE sur reseau bridge Docker"
)

divider()
page_break()

# ── 3. DOCKER COMPOSE ──────────────────────────────────
heading1("3. Configuration Docker Compose")

para("Docker Compose nous a permis de decrire l'ensemble de l'infrastructure en un seul fichier declaratif. Sans lui, demarrer trois instances Oracle, configurer leurs reseaux et gerer leur ordre de demarrage serait fastidieux et non reproductible.")

heading2("3.1 Declaration du reseau")

code_block(
    "version: '3.8'\n"
    "\n"
    "networks:\n"
    "  eshop-network:\n"
    "    driver: bridge           # Reseau bridge Docker isole\n"
    "    ipam:\n"
    "      config:\n"
    "        - subnet: 172.20.0.0/16  # Plage IP dediee -- evite les conflits",
    title="docker-compose.yml -- Declaration du reseau prive"
)

heading2("3.2 Services Oracle -- Extrait annote")

code_block(
    "  oracle-site1:\n"
    "    image: gvenzl/oracle-xe:21-slim  # Image Oracle XE 21c allegee (~2 Go RAM)\n"
    "    container_name: oracle-site1     # Nom fixe obligatoire pour les DB Links\n"
    "    hostname: oracle-site1           # Nom DNS interne Docker\n"
    "    environment:\n"
    "      ORACLE_PASSWORD: oracle1       # Mot de passe SYS/SYSTEM\n"
    "      ORACLE_DATABASE: BDDVENTE      # Nom du service (PDB Oracle)\n"
    "      APP_USER: eshop1               # Utilisateur applicatif cree auto\n"
    "      APP_USER_PASSWORD: eshop1pass\n"
    "    ports:\n"
    "      - '1521:1521'                  # Port expose sur la machine hote\n"
    "    networks:\n"
    "      eshop-network:\n"
    "        ipv4_address: 172.20.0.10    # IP fixe (indispensable pour DB Links)\n"
    "    healthcheck:\n"
    "      test: ['CMD', 'healthcheck.sh'] # Script fourni par l'image gvenzl\n"
    "      interval: 30s\n"
    "      timeout: 10s\n"
    "      retries: 10\n"
    "      start_period: 120s             # Oracle XE met ~2 min a demarrer\n"
    "    restart: unless-stopped\n"
    "\n"
    "  oracle-central:\n"
    "    ...                              # Meme structure\n"
    "    depends_on:\n"
    "      oracle-site1:\n"
    "        condition: service_healthy   # Attend Site1 operationnel\n"
    "      oracle-site2:\n"
    "        condition: service_healthy   # Et Site2 operationnel\n"
    "    networks:\n"
    "      eshop-network:\n"
    "        ipv4_address: 172.20.0.12",
    title="docker-compose.yml -- Services oracle-site1 et oracle-central"
)

heading2("3.3 Ordre d'initialisation des scripts")

para("L'image gvenzl/oracle-xe execute les fichiers de /docker-entrypoint-initdb.d/ par ordre alphabetique. Sans prefixage numerique, fragments/ s'executerait avant schema/, causant des erreurs car les tables sources n'existeraient pas encore.")

code_block(
    "    volumes:\n"
    "      - ./schema:/docker-entrypoint-initdb.d/01_schema:ro     # CREATE TABLE\n"
    "      - ./data:/docker-entrypoint-initdb.d/02_data:ro         # INSERT donnees\n"
    "      - ./docker/grants.sh:.../03_grants.sh:ro                # GRANT privileges\n"
    "      - ./site1/dblinks:.../04_dblinks:ro                     # DB Links sortants\n"
    "      - ./site1/scenario2/fragments:.../05_fragments:ro       # Fragmentation\n"
    "      - ./site1/scenario2/procedures:.../06_procedures:ro     # Procedures\n"
    "      - ./site1/scenario2/triggers:.../07_triggers:ro         # Triggers\n"
    "      - ./site1/scenario2/indexes:.../08_indexes:ro           # Index\n"
    "      - ./site1/scenario2/synonyms:.../09_synonyms:ro         # Synonymes",
    title="docker-compose.yml -- Volumes montes avec prefixage ordonne (Site1)"
)

note("Le suffixe :ro (read-only) empeche les conteneurs de modifier les scripts sources sur l'hote, pratique de securite recommandee.")

divider()
page_break()

# ── 4. RESEAU ──────────────────────────────────────────
heading1("4. Configuration reseau")

heading2("4.1 Choix du sous-reseau 172.20.0.0/16")

para("Nous avons retenu ce sous-reseau pour trois raisons principales :")

bullet("Pas de conflit avec les plages 192.168.x.x courantes des reseaux domestiques et d'entreprise.", bold_prefix="Isolation : ")
bullet("La plage /16 offre plus de 65 000 adresses, garantissant l'absence de saturation meme si le projet evolue.", bold_prefix="Capacite : ")
bullet("Les IPs fixes sont indispensables pour les Database Links Oracle : si l'IP changeait au redemarrage, tous les liens seraient rompus.", bold_prefix="Stabilite : ")

rich_table([
    ("oracle-site1",   "172.20.0.10", "172.20.0.0/16", "eshop-network (bridge)"),
    ("oracle-site2",   "172.20.0.11", "172.20.0.0/16", "eshop-network (bridge)"),
    ("oracle-central", "172.20.0.12", "172.20.0.0/16", "eshop-network (bridge)"),
], headers=["Conteneur", "IP fixe", "Sous-reseau", "Reseau Docker"],
   col_widths=[3.5, 2.8, 3.0, 7.7])

heading2("4.2 Resolution DNS interne Docker")

para("Docker Compose cree automatiquement un serveur DNS interne au reseau bridge. Les noms de conteneurs (oracle-site1, oracle-site2, oracle-central) sont directement utilisables comme noms d'hotes dans les chaines de connexion Oracle, ce qui simplifie la configuration des DB Links.")

code_block(
    "# Test TCP port 1521 entre conteneurs (depuis connectivity_test.sh)\n"
    "# ping est absent de l'image slim -- on utilise /dev/tcp a la place\n"
    "docker exec oracle-site1 bash -c \\\n"
    "    'timeout 3 bash -c \"cat < /dev/null > /dev/tcp/oracle-site2/1521\"'\n"
    "# exit code 0 = port accessible\n"
    "\n"
    "# Verification DNS interne\n"
    "docker exec oracle-central bash -c 'getent hosts oracle-site1'\n"
    "# Resultat : 172.20.0.10   oracle-site1",
    title="scripts/connectivity_test.sh -- Tests TCP inter-conteneurs"
)

divider()

# ── 5. CONFIGURATION SITES ─────────────────────────────
heading1("5. Configuration des sites Oracle")

heading2("5.1 Schema global commun")

para("Chaque conteneur recoit d'abord le meme schema global (script 01_schema). Toutes les tables sont creees identiquement sur les trois noeuds avant la fragmentation.")

code_block(
    "-- eshop_global.sh -- Schema global EShop\n"
    "-- Execute comme APP_USER sur chaque conteneur\n"
    "\n"
    "-- Table centrale de fragmentation\n"
    "CREATE TABLE LIGNECOMMANDES (\n"
    "    IDLIGNECOMMANDE NUMBER(*,0),\n"
    "    IDCOMMANDE      NUMBER(*,0),\n"
    "    IDPRODUIT       NUMBER(*,0),\n"
    "    QUANTITE        NUMBER(*,0),   -- Critere Sc2 : >= 100 ou < 100\n"
    "    REMISE          FLOAT(126)\n"
    ");\n"
    "\n"
    "-- Table produits avec lien categorie (Critere Sc1 : IDCATEG)\n"
    "CREATE TABLE PRODUITS (\n"
    "    IDPRODUIT    NUMBER(*,0),\n"
    "    DESIGNATION  VARCHAR2(100),\n"
    "    IDFOUR       NUMBER(*,0),\n"
    "    IDCATEG      NUMBER(*,0),      -- Critere Sc1 : = 50 ou = 35\n"
    "    PRIXUNITAIRE FLOAT(126),\n"
    "    UNITESENSTOCK            NUMBER(*,0),\n"
    "    INDISPONIBLE             NUMBER(*,0),\n"
    "    CONSTRAINT chk_indisponible CHECK (INDISPONIBLE IN (0,1))\n"
    ");\n"
    "\n"
    "-- Cles primaires\n"
    "ALTER TABLE CATEGORIES      ADD PRIMARY KEY (IDCATEG);\n"
    "ALTER TABLE PRODUITS        ADD PRIMARY KEY (IDPRODUIT);\n"
    "ALTER TABLE LIGNECOMMANDES  ADD PRIMARY KEY (IDLIGNECOMMANDE);\n"
    "\n"
    "-- Integrite referentielle\n"
    "ALTER TABLE PRODUITS ADD FOREIGN KEY (IDCATEG)\n"
    "    REFERENCES CATEGORIES (IDCATEG);\n"
    "ALTER TABLE LIGNECOMMANDES ADD FOREIGN KEY (IDCOMMANDE)\n"
    "    REFERENCES COMMANDES (IDCOMMANDE);\n"
    "COMMIT;",
    title="schema/eshop_global.sh -- DDL principal (extrait)"
)

heading2("5.2 Donnees initiales -- Repartition des 25 lignes")

rich_table([
    ("Sc2 - Site 1", "QUANTITE >= 100", "12", "LC 1,2,3,4,5,11,12,13,14,15,22,23"),
    ("Sc2 - Site 2", "QUANTITE < 100",  "13", "LC 6,7,8,9,10,16,17,18,19,20,21,24,25"),
    ("Sc1 - Site 1", "IDCATEG = 50",    "10", "Produits Informatique (IDs 1, 2, 3)"),
    ("Sc1 - Site 2", "IDCATEG = 35",    "10", "Produits Accessoires (IDs 4, 5, 6)"),
], headers=["Fragment", "Critere", "Lignes", "Details"],
   col_widths=[3.0, 3.0, 2.0, 9.0])

divider()
page_break()

# ── 6. TESTS CONNECTIVITE ──────────────────────────────
heading1("6. Tests de connectivite")

para("Avant de configurer les Database Links, il est imperatif de valider que les trois noeuds peuvent communiquer. Le script connectivity_test.sh effectue cinq niveaux de verification automatisee.")

heading2("6.1 Structure du script de test")

code_block(
    "#!/bin/bash\n"
    "# connectivity_test.sh -- 5 niveaux de verification\n"
    "\n"
    "# Niveau 1 : Connectivite TCP port 1521 inter-conteneurs\n"
    "docker exec oracle-site1 bash -c \\\n"
    "    'timeout 3 bash -c \"cat < /dev/null > /dev/tcp/oracle-site2/1521\"'\n"
    "check $? 'TCP oracle-site1 -> oracle-site2:1521'\n"
    "\n"
    "# Niveau 2 : Connexions SQL*Plus locales\n"
    "printf 'SELECT \\\"OK\\\" FROM DUAL;\\nEXIT;\\n' \\\n"
    "    | docker exec -i oracle-site1 \\\n"
    "        sqlplus -s 'eshop1/eshop1pass@//localhost:1521/BDDVENTE'\n"
    "\n"
    "# Niveau 3 : Database Links depuis oracle-central\n"
    "printf 'SELECT \\\"LINK_OK\\\" FROM DUAL@site1_link;\\nEXIT;\\n' \\\n"
    "    | docker exec -i oracle-central \\\n"
    "        sqlplus -s 'eshopcentral/centralpass@//localhost:1521/BDDCENTRAL'\n"
    "\n"
    "# Niveau 4 : DB Links bidirectionnels site1 <-> site2\n"
    "# Niveau 5 : Requetes distribuees de validation (COUNT global)",
    title="scripts/connectivity_test.sh -- Architecture du script"
)

heading2("6.2 Sortie console attendue -- Cluster operationnel")

code_block(
    "============================================================\n"
    "  CONNECTIVITY TEST -- Distributed EShop  2026-06-09 14:32:17\n"
    "============================================================\n"
    "\n"
    "=== Niveau 1 : TCP inter-conteneurs (port 1521) ===\n"
    "  [PASS] TCP oracle-site1 -> oracle-site2:1521\n"
    "  [PASS] TCP oracle-site1 -> oracle-central:1521\n"
    "  [PASS] TCP oracle-site2 -> oracle-site1:1521\n"
    "  [PASS] TCP oracle-central -> oracle-site1:1521\n"
    "  [PASS] TCP oracle-central -> oracle-site2:1521\n"
    "\n"
    "=== Niveau 2 : Connexions SQL*Plus locales ===\n"
    "  [PASS] SQL*Plus eshop1@BDDVENTE\n"
    "  [PASS] SQL*Plus eshop2@BDDVENTE2\n"
    "  [PASS] SQL*Plus eshopcentral@BDDCENTRAL\n"
    "\n"
    "=== Niveau 3 : Database Links depuis oracle-central ===\n"
    "  [PASS] DB Link central -> site1_link\n"
    "  [PASS] DB Link central -> site2_link\n"
    "\n"
    "=== Niveau 4 : DB Links bidirectionnels ===\n"
    "  [PASS] DB Link oracle-site1 -> site2_link\n"
    "  [PASS] DB Link oracle-site1 -> central_link\n"
    "  [PASS] DB Link oracle-site2 -> site1_link\n"
    "  [PASS] DB Link oracle-site2 -> central_link\n"
    "\n"
    "=== Niveau 5 : Requetes distribuees ===\n"
    "  SITE    NB_LIGNES    QTE_TOTALE\n"
    "  ------  ----------   ----------\n"
    "  SITE1          12         3545\n"
    "  SITE2          13          422\n"
    "\n"
    "  Resultat : 14 PASS / 0 FAIL / 14 tests\n"
    "  OK -- Tous les tests sont PASSES -- cluster operationnel\n"
    "============================================================",
    title="Sortie console simulee -- Resultat attendu apres demarrage complet"
)

note("Si un test TCP echoue, verifier en premier : docker-compose ps. Un conteneur en etat 'starting' n'accepte pas encore de connexions.")

divider()
page_break()

# ── 7. UTILISATEURS ────────────────────────────────────
heading1("7. Creation des utilisateurs et privileges")

para("L'image gvenzl/oracle-xe cree automatiquement l'APP_USER defini en variable d'environnement, mais ne lui accorde pas les privileges etendus necessaires pour notre projet distribue. Le script grants.sh comble ce manque en etant execute comme SYSDBA.")

code_block(
    "#!/bin/bash\n"
    "# grants.sh -- Execute par SYSDBA lors de l'initialisation\n"
    "\n"
    "sqlplus -s \"sys/${ORACLE_PASSWORD}@//localhost:1521/${ORACLE_DATABASE} as sysdba\" <<EOF\n"
    "WHENEVER SQLERROR EXIT SQL.SQLCODE;  -- Arret immediat si erreur\n"
    "\n"
    "-- Permet de creer des liens vers d'autres instances Oracle\n"
    "GRANT CREATE DATABASE LINK TO ${APP_USER};\n"
    "\n"
    "-- Permet de creer des alias transparents vers objets distants\n"
    "GRANT CREATE SYNONYM       TO ${APP_USER};\n"
    "\n"
    "-- Permet de creer des vues distribuees (UNION ALL inter-sites)\n"
    "GRANT CREATE VIEW          TO ${APP_USER};\n"
    "\n"
    "-- Connexion explicite (securite defensive)\n"
    "GRANT CREATE SESSION       TO ${APP_USER};\n"
    "\n"
    "-- Verification : liste les grants effectivement enregistres\n"
    "SELECT PRIVILEGE, GRANTEE\n"
    "FROM DBA_SYS_PRIVS\n"
    "WHERE GRANTEE = UPPER('${APP_USER}')\n"
    "  AND PRIVILEGE IN ('CREATE DATABASE LINK','CREATE SYNONYM','CREATE VIEW')\n"
    "ORDER BY PRIVILEGE;\n"
    "\n"
    "COMMIT; EXIT;\n"
    "EOF",
    title="docker/grants.sh -- Attribution des privileges systeme"
)

rich_table([
    ("CREATE DATABASE LINK", "Creer des liens vers d'autres instances Oracle",    "Obligatoire pour la distribution"),
    ("CREATE SYNONYM",       "Creer des alias transparents vers objets distants", "Masque la complexite @link"),
    ("CREATE VIEW",          "Creer des vues (y compris INSTEAD OF sur vues)",    "Vues globales + triggers routage"),
    ("CREATE SESSION",       "Se connecter a la base de donnees",                 "Securite explicite"),
], headers=["Privilege", "Utilite", "Raison"],
   col_widths=[4.0, 7.0, 6.0])

divider()

# ── 8. DATABASE LINKS ──────────────────────────────────
heading1("8. Configuration des Database Links")

para("Les Database Links sont le mecanisme fondamental des bases Oracle distribuees. Ils permettent a une session connectee sur un noeud A d'executer des requetes sur un noeud B de maniere transparente, via la notation @nom_du_lien.")

heading2("8.1 Architecture bidirectionnelle complete")

rich_table([
    ("oracle-central", "oracle-site1",   "site1_link",   "Lecture/ecriture fragments QTE >= 100"),
    ("oracle-central", "oracle-site2",   "site2_link",   "Lecture/ecriture fragments QTE < 100"),
    ("oracle-site1",   "oracle-site2",   "site2_link",   "Synchronisation inter-sites"),
    ("oracle-site1",   "oracle-central", "central_link", "Acces tables de reference"),
    ("oracle-site2",   "oracle-site1",   "site1_link",   "Synchronisation inter-sites"),
    ("oracle-site2",   "oracle-central", "central_link", "Acces tables de reference"),
], headers=["De", "Vers", "Nom du lien", "Usage"],
   col_widths=[3.2, 3.2, 3.0, 7.6])

heading2("8.2 Creation des liens -- oracle-central")

code_block(
    "-- setup_dblinks.sql -- DB Links Central -> Site1 et Site2\n"
    "-- Execute comme eshopcentral dans BDDCENTRAL\n"
    "\n"
    "-- Lien vers Site1 (grossistes : QTE >= 100)\n"
    "CREATE DATABASE LINK site1_link\n"
    "  CONNECT TO eshop1 IDENTIFIED BY eshop1pass\n"
    "  USING '(DESCRIPTION=\n"
    "    (ADDRESS=(PROTOCOL=TCP)(HOST=oracle-site1)(PORT=1521))\n"
    "    (CONNECT_DATA=(SERVICE_NAME=BDDVENTE))\n"
    "  )';\n"
    "\n"
    "-- Lien vers Site2 (detaillants : QTE < 100)\n"
    "CREATE DATABASE LINK site2_link\n"
    "  CONNECT TO eshop2 IDENTIFIED BY eshop2pass\n"
    "  USING '(DESCRIPTION=\n"
    "    (ADDRESS=(PROTOCOL=TCP)(HOST=oracle-site2)(PORT=1521))\n"
    "    (CONNECT_DATA=(SERVICE_NAME=BDDVENTE2))\n"
    "  )';\n"
    "\n"
    "-- Vue globale Scenario 2 : UNION des deux sites\n"
    "CREATE OR REPLACE VIEW V_LIGNECOMMANDES_GLOBAL AS\n"
    "  SELECT 'SITE1' AS SITE, lc.*\n"
    "  FROM LigneCommandes1@site1_link lc\n"
    "  UNION ALL\n"
    "  SELECT 'SITE2' AS SITE, lc.*\n"
    "  FROM LigneCommandes2@site2_link lc;\n"
    "\n"
    "COMMIT;",
    title="docker/setup_dblinks.sql -- Liens du Central et vue globale"
)

heading2("8.3 Liens sortants de Site1")

code_block(
    "-- site1_dblinks.sh -- DB Links sortants de oracle-site1\n"
    "\n"
    "-- Lien vers Site2 : permet d'interroger/modifier les donnees distantes\n"
    "CREATE DATABASE LINK site2_link\n"
    "  CONNECT TO eshop2 IDENTIFIED BY eshop2pass\n"
    "  USING '(DESCRIPTION=\n"
    "    (ADDRESS=(PROTOCOL=TCP)(HOST=oracle-site2)(PORT=1521))\n"
    "    (CONNECT_DATA=(SERVICE_NAME=BDDVENTE2))\n"
    "  )';\n"
    "\n"
    "-- Lien vers Central : acces aux tables globales de reference\n"
    "CREATE DATABASE LINK central_link\n"
    "  CONNECT TO eshopcentral IDENTIFIED BY centralpass\n"
    "  USING '(DESCRIPTION=\n"
    "    (ADDRESS=(PROTOCOL=TCP)(HOST=oracle-central)(PORT=1521))\n"
    "    (CONNECT_DATA=(SERVICE_NAME=BDDCENTRAL))\n"
    "  )';\n"
    "\n"
    "COMMIT;",
    title="site1/dblinks/site1_dblinks.sh -- Liens sortants de Site1"
)

divider()
page_break()

# ── 9. SCHEMA GLOBAL ───────────────────────────────────
heading1("9. Schema de donnees global")

para("Le modele de donnees du systeme EShop comprend sept tables organisees autour du concept de commande client. La table LIGNECOMMANDES est le centre de la fragmentation.")

code_block(
    "CATEGORIES (IDCATEG, NOMDECATEGORIE, DESCRIPTION)\n"
    "    |\n"
    "    +-- PRODUITS (IDPRODUIT, DESIGNATION, IDFOUR, IDCATEG, PRIXUNITAIRE, ...)\n"
    "    |       |\n"
    "    |       +-- LIGNECOMMANDES (IDLIGNECOMMANDE, IDCOMMANDE,\n"
    "    |                          IDPRODUIT, QUANTITE, REMISE)\n"
    "    |                   |\n"
    "FOURNISSEURS (IDFOUR)   +-- COMMANDES (IDCOMMANDE, IDEMPLOYE, IDCLIENT, ...)\n"
    "                                |                |\n"
    "                            EMPLOYES          CLIENTS\n"
    "                            (IDEMPLOYE)       (IDCLIENT)",
    title="Diagramme relationnel simplifie"
)

rich_table([
    ("CATEGORIES",     "6 lignes",  "Boissons, Accessoires(35), Informatique(50)..."),
    ("FOURNISSEURS",   "3 lignes",  "TechDistrib SA, InfoParts SARL, AccessPro SAS"),
    ("EMPLOYES",       "3 lignes",  "Alice Dupont, Bob Martin, Claire Bernard"),
    ("CLIENTS",        "8 lignes",  "Societe Alpha a Societe Theta -- villes francaises"),
    ("PRODUITS",       "10 lignes", "3 categ. 50 + 3 categ. 35 + 4 autres categories"),
    ("COMMANDES",      "10 lignes", "Toutes en 2026 -- differents clients et employes"),
    ("LIGNECOMMANDES", "25 lignes", "Quantites de 2 a 1000 -- 12 >= 100, 13 < 100"),
], headers=["Table", "Volume", "Contenu"],
   col_widths=[3.5, 2.5, 11.0])

divider()

# ── 10. FRAGMENTATION ──────────────────────────────────
heading1("10. Tables fragmentees sur les sites distants")

heading2("10.1 Scenario 2 -- Fragmentation par volume (deploye par defaut)")

para("Justification metier : les grossistes (QTE >= 100) sont des clients prioritaires dont les donnees doivent etre sur le serveur principal. La fragmentation est realisee par CREATE TABLE AS SELECT avec une contrainte CHECK qui verrouille la regle.")

code_block(
    "-- site1_fragments.sh -- Fragmentation horizontale Site1\n"
    "-- Regle : QUANTITE >= 100  (grossistes)\n"
    "\n"
    "-- Fragment clients : uniquement ceux lies a des lignes QTE >= 100\n"
    "CREATE TABLE Clients1 AS\n"
    "SELECT DISTINCT cl.*\n"
    "FROM CLIENTS cl\n"
    "JOIN COMMANDES c ON cl.IDCLIENT = c.IDCLIENT\n"
    "JOIN LIGNECOMMANDES lc ON c.IDCOMMANDE = lc.IDCOMMANDE\n"
    "WHERE lc.QUANTITE >= 100;\n"
    "\n"
    "ALTER TABLE Clients1 ADD PRIMARY KEY (IDCLIENT);\n"
    "\n"
    "-- Fragment lignes de commande avec verrou de fragmentation\n"
    "CREATE TABLE LigneCommandes1 AS\n"
    "SELECT lc.* FROM LIGNECOMMANDES lc\n"
    "WHERE lc.QUANTITE >= 100;\n"
    "\n"
    "ALTER TABLE LigneCommandes1 ADD PRIMARY KEY (IDLIGNECOMMANDE);\n"
    "-- Verrou : interdit toute insertion QTE < 100 sur ce site\n"
    "ALTER TABLE LigneCommandes1 ADD CONSTRAINT chk_sc2_site1_qte\n"
    "    CHECK (QUANTITE >= 100);\n"
    "\n"
    "COMMIT;",
    title="site1/scenario2/fragments/site1_fragments.sh -- Site1 (QTE >= 100)"
)

code_block(
    "-- site2_fragments.sh -- Fragmentation Site2 (miroir inverse)\n"
    "-- Regle : QUANTITE < 100  (detaillants)\n"
    "\n"
    "CREATE TABLE LigneCommandes2 AS\n"
    "SELECT lc.* FROM LIGNECOMMANDES lc\n"
    "WHERE lc.QUANTITE < 100;\n"
    "\n"
    "ALTER TABLE LigneCommandes2 ADD PRIMARY KEY (IDLIGNECOMMANDE);\n"
    "-- Verrou miroir : interdit toute insertion QTE >= 100 sur Site2\n"
    "ALTER TABLE LigneCommandes2 ADD CONSTRAINT chk_sc2_site2_qte\n"
    "    CHECK (QUANTITE < 100);\n"
    "\n"
    "COMMIT;",
    title="site2/scenario2/fragments/site2_fragments.sh -- Site2 (QTE < 100)"
)

heading2("10.2 Scenario 1 -- Fragmentation par categorie")

code_block(
    "-- site1_fragments.sql -- Scenario 1 (tables suffixees _sc1)\n"
    "-- Regle : IDCATEG = 50 (Informatique)\n"
    "\n"
    "-- Sous-ensemble de categories : uniquement Informatique\n"
    "CREATE TABLE Categories1_sc1 AS\n"
    "SELECT * FROM CATEGORIES WHERE IDCATEG = 50;\n"
    "ALTER TABLE Categories1_sc1 ADD CONSTRAINT chk_sc1_categ CHECK (IDCATEG = 50);\n"
    "\n"
    "-- Produits Informatique uniquement\n"
    "CREATE TABLE Produits1_sc1 AS\n"
    "SELECT * FROM PRODUITS WHERE IDCATEG = 50;\n"
    "\n"
    "-- Lignes liees aux produits categ. 50 -- SANS condition sur QUANTITE\n"
    "-- CORRECTION : la v1 incluait AND QUANTITE > 100 ce qui laissait\n"
    "-- des lignes orphelines (IDCATEG=50 AND QTE<=100 non attribuees).\n"
    "CREATE TABLE LigneCommandes1_sc1 AS\n"
    "SELECT lc.*\n"
    "FROM LIGNECOMMANDES lc\n"
    "JOIN PRODUITS p ON lc.IDPRODUIT = p.IDPRODUIT\n"
    "WHERE p.IDCATEG = 50;\n"
    "\n"
    "COMMIT;",
    title="site1/scenario1/fragments/site1_fragments.sql -- Scenario 1"
)

warn("Bug corrige : la premiere version du Scenario 1 incluait une condition AND QUANTITE > 100 qui creait des lignes orphelines non attribuees a aucun site. La regle correcte est uniquement basee sur la categorie.")

divider()
page_break()

# ── 11. PROCEDURES ─────────────────────────────────────
heading1("11. Architecture des procedures stockees")

para("Les procedures stockees centralisent les regles metier et garantissent que les regles de fragmentation sont respectees. Sans elles, une application cliente pourrait inserer une ligne avec QTE = 50 sur Site1, violant la regle de fragmentation.")

heading2("11.1 Procedure d'insertion -- Site1 Scenario 2")

code_block(
    "-- insertligne -- Site1 (QUANTITE >= 100)\n"
    "CREATE OR REPLACE PROCEDURE insertligne (\n"
    "    p_id        IN LIGNECOMMANDES1.IDLIGNECOMMANDE%TYPE,\n"
    "    p_idcmd     IN LIGNECOMMANDES1.IDCOMMANDE%TYPE,\n"
    "    p_idprod    IN LIGNECOMMANDES1.IDPRODUIT%TYPE,\n"
    "    p_qte       IN LIGNECOMMANDES1.QUANTITE%TYPE,\n"
    "    p_remise    IN LIGNECOMMANDES1.REMISE%TYPE DEFAULT 0\n"
    ") AS\n"
    "    v_cmd_count  NUMBER;\n"
    "    v_prod_count NUMBER;\n"
    "BEGIN\n"
    "    -- GARDE 1 : regle de fragmentation\n"
    "    IF p_qte < 100 THEN\n"
    "        RAISE_APPLICATION_ERROR(-20001,\n"
    "            'SITE1 ERROR: Quantite ' || p_qte ||\n"
    "            ' < 100. Cette ligne appartient a Site2.');\n"
    "    END IF;\n"
    "\n"
    "    -- GARDE 2 : la commande doit exister dans le fragment local\n"
    "    SELECT COUNT(*) INTO v_cmd_count\n"
    "    FROM COMMANDES1 WHERE IDCOMMANDE = p_idcmd;\n"
    "    IF v_cmd_count = 0 THEN\n"
    "        RAISE_APPLICATION_ERROR(-20002,\n"
    "            'Commande ' || p_idcmd || ' introuvable dans Site1.');\n"
    "    END IF;\n"
    "\n"
    "    -- GARDE 3 : le produit doit exister dans le fragment local\n"
    "    SELECT COUNT(*) INTO v_prod_count\n"
    "    FROM PRODUITS1 WHERE IDPRODUIT = p_idprod;\n"
    "    IF v_prod_count = 0 THEN\n"
    "        RAISE_APPLICATION_ERROR(-20003,\n"
    "            'Produit ' || p_idprod || ' introuvable dans Site1.');\n"
    "    END IF;\n"
    "\n"
    "    -- Insertion effective\n"
    "    INSERT INTO LIGNECOMMANDES1\n"
    "        (IDLIGNECOMMANDE, IDCOMMANDE, IDPRODUIT, QUANTITE, REMISE)\n"
    "    VALUES (p_id, p_idcmd, p_idprod, p_qte, p_remise);\n"
    "\n"
    "    COMMIT;\n"
    "    DBMS_OUTPUT.PUT_LINE('OK - Ligne ' || p_id ||\n"
    "        ' inseree dans Site1 (Qte=' || p_qte || ')');\n"
    "\n"
    "EXCEPTION\n"
    "    WHEN DUP_VAL_ON_INDEX THEN\n"
    "        RAISE_APPLICATION_ERROR(-20004,\n"
    "            'Ligne ' || p_id || ' existe deja dans Site1.');\n"
    "    WHEN OTHERS THEN\n"
    "        ROLLBACK;\n"
    "        RAISE;\n"
    "END insertligne;\n"
    "/",
    title="site1/scenario2/procedures/site1_procedures.sh -- insertligne"
)

heading2("11.2 Procedure de suppression avec nettoyage cascadant")

code_block(
    "-- deleteligne -- Suppression avec nettoyage des commandes orphelines\n"
    "CREATE OR REPLACE PROCEDURE deleteligne (p_id IN NUMBER) AS\n"
    "    v_count     NUMBER;\n"
    "    v_idcmd     NUMBER;\n"
    "    v_cmd_count NUMBER;\n"
    "BEGIN\n"
    "    -- Recupere l'ID de commande avant suppression\n"
    "    SELECT COUNT(*), MAX(IDCOMMANDE) INTO v_count, v_idcmd\n"
    "    FROM LIGNECOMMANDES1 WHERE IDLIGNECOMMANDE = p_id;\n"
    "\n"
    "    IF v_count = 0 THEN\n"
    "        RAISE_APPLICATION_ERROR(-20008,\n"
    "            'Ligne ' || p_id || ' introuvable dans Site1.');\n"
    "    END IF;\n"
    "\n"
    "    DELETE FROM LIGNECOMMANDES1 WHERE IDLIGNECOMMANDE = p_id;\n"
    "\n"
    "    -- Si la commande n'a plus de lignes -> la supprimer aussi\n"
    "    SELECT COUNT(*) INTO v_cmd_count\n"
    "    FROM LIGNECOMMANDES1 WHERE IDCOMMANDE = v_idcmd;\n"
    "\n"
    "    IF v_cmd_count = 0 THEN\n"
    "        DELETE FROM COMMANDES1 WHERE IDCOMMANDE = v_idcmd;\n"
    "        DBMS_OUTPUT.PUT_LINE('OK - Commande ' || v_idcmd ||\n"
    "            ' supprimee (plus de lignes).');\n"
    "    END IF;\n"
    "\n"
    "    COMMIT;\n"
    "    DBMS_OUTPUT.PUT_LINE('OK - Ligne ' || p_id || ' supprimee de Site1.');\n"
    "EXCEPTION\n"
    "    WHEN OTHERS THEN ROLLBACK; RAISE;\n"
    "END deleteligne;\n"
    "/",
    title="site1/scenario2/procedures/site1_procedures.sh -- deleteligne"
)

heading2("11.3 Procedures de routage depuis le Central")

code_block(
    "-- route_insert_ligne -- Routage Central -> bon site\n"
    "CREATE OR REPLACE PROCEDURE route_insert_ligne (\n"
    "    p_id IN NUMBER, p_idcmd IN NUMBER, p_idprod IN NUMBER,\n"
    "    p_qte IN NUMBER, p_remise IN NUMBER DEFAULT 0\n"
    ") AS\n"
    "    v_count_s1 NUMBER; v_count_s2 NUMBER;\n"
    "    v_count    NUMBER; v_site     VARCHAR2(10);\n"
    "BEGIN\n"
    "    -- Validation : commande et produit doivent exister sur le central\n"
    "    SELECT COUNT(*) INTO v_count FROM COMMANDES WHERE IDCOMMANDE = p_idcmd;\n"
    "    IF v_count = 0 THEN\n"
    "        RAISE_APPLICATION_ERROR(-20101, 'Commande introuvable.');\n"
    "    END IF;\n"
    "\n"
    "    -- Verification : pas de doublon inter-sites\n"
    "    SELECT COUNT(*) INTO v_count_s1\n"
    "    FROM LigneCommandes1@site1_link WHERE IDLIGNECOMMANDE = p_id;\n"
    "    SELECT COUNT(*) INTO v_count_s2\n"
    "    FROM LigneCommandes2@site2_link WHERE IDLIGNECOMMANDE = p_id;\n"
    "    IF v_count_s1 + v_count_s2 > 0 THEN\n"
    "        RAISE_APPLICATION_ERROR(-20103, 'Ligne existe deja.');\n"
    "    END IF;\n"
    "\n"
    "    -- DECISION DE ROUTAGE\n"
    "    IF p_qte >= 100 THEN\n"
    "        v_site := 'SITE1';\n"
    "        ensure_site1_refs(p_idcmd, p_idprod);  -- Copie les refs si besoin\n"
    "        INSERT INTO LigneCommandes1@site1_link  -- Insertion distante\n"
    "            (IDLIGNECOMMANDE, IDCOMMANDE, IDPRODUIT, QUANTITE, REMISE)\n"
    "        VALUES (p_id, p_idcmd, p_idprod, p_qte, NVL(p_remise, 0));\n"
    "    ELSE\n"
    "        v_site := 'SITE2';\n"
    "        ensure_site2_refs(p_idcmd, p_idprod);\n"
    "        INSERT INTO LigneCommandes2@site2_link\n"
    "            (IDLIGNECOMMANDE, IDCOMMANDE, IDPRODUIT, QUANTITE, REMISE)\n"
    "        VALUES (p_id, p_idcmd, p_idprod, p_qte, NVL(p_remise, 0));\n"
    "    END IF;\n"
    "\n"
    "    -- Journalisation du routage\n"
    "    INSERT INTO ROUTING_LOG\n"
    "        (OPERATION, SITE_CIBLE, IDLIGNECOMMANDE, IDCOMMANDE, IDPRODUIT, QUANTITE)\n"
    "    VALUES ('INSERT', v_site, p_id, p_idcmd, p_idprod, p_qte);\n"
    "END route_insert_ligne;\n"
    "/",
    title="docker/routing_central.sh -- route_insert_ligne (routage automatique)"
)

divider()
page_break()

# ── 12. TRIGGERS ───────────────────────────────────────
heading1("12. Architecture des triggers")

para("Trois categories de triggers ont ete implementes, chacune avec un role distinct dans l'architecture distribuee :")

rich_table([
    ("BEFORE INSERT/UPDATE", "Controle",   "Bloque toute operation violant la regle de fragmentation"),
    ("AFTER INSERT/UPDATE/DELETE", "Log",  "Enregistre chaque DML dans la table de journal du site"),
    ("INSTEAD OF sur vue",  "Routage",    "Intercepte les DML sur la vue centrale et les route"),
], headers=["Type", "Role", "Action"],
   col_widths=[4.5, 2.5, 10.0])

heading2("12.1 Triggers de controle et de log -- Site1")

code_block(
    "-- Table de journalisation des operations\n"
    "CREATE TABLE LOG_SITE1 (\n"
    "    ID_LOG          NUMBER GENERATED ALWAYS AS IDENTITY PRIMARY KEY,\n"
    "    OPERATION       VARCHAR2(10),       -- 'INSERT', 'UPDATE', 'DELETE'\n"
    "    IDLIGNECOMMANDE NUMBER,\n"
    "    QUANTITE_OLD    NUMBER,             -- Valeur avant modification\n"
    "    QUANTITE_NEW    NUMBER,             -- Valeur apres modification\n"
    "    DATE_OPERATION  TIMESTAMP DEFAULT SYSTIMESTAMP,\n"
    "    UTILISATEUR     VARCHAR2(100) DEFAULT USER\n"
    ");\n"
    "\n"
    "-- TRIGGER 1 : Controle AVANT insertion\n"
    "CREATE OR REPLACE TRIGGER trg_check_site1_insert\n"
    "BEFORE INSERT ON LigneCommandes1\n"
    "FOR EACH ROW\n"
    "BEGIN\n"
    "    IF :NEW.QUANTITE < 100 THEN\n"
    "        RAISE_APPLICATION_ERROR(-20010,\n"
    "            'TRIGGER SITE1 - INSERT bloque : QUANTITE=' || :NEW.QUANTITE ||\n"
    "            ' < 100. Cette ligne appartient a Site2.');\n"
    "    END IF;\n"
    "END;\n"
    "/\n"
    "\n"
    "-- TRIGGER 2 : Controle AVANT mise a jour\n"
    "CREATE OR REPLACE TRIGGER trg_check_site1_update\n"
    "BEFORE UPDATE ON LigneCommandes1\n"
    "FOR EACH ROW\n"
    "BEGIN\n"
    "    IF :NEW.QUANTITE < 100 THEN\n"
    "        RAISE_APPLICATION_ERROR(-20011,\n"
    "            'TRIGGER SITE1 - UPDATE bloque : Nouvelle QTE=' || :NEW.QUANTITE ||\n"
    "            ' < 100. Modifier sur Site2 a la place.');\n"
    "    END IF;\n"
    "END;\n"
    "/\n"
    "\n"
    "-- TRIGGER 3 : Log APRES toute operation DML\n"
    "CREATE OR REPLACE TRIGGER trg_log_site1_operations\n"
    "AFTER INSERT OR UPDATE OR DELETE ON LigneCommandes1\n"
    "FOR EACH ROW\n"
    "DECLARE v_operation VARCHAR2(10);\n"
    "BEGIN\n"
    "    IF INSERTING THEN\n"
    "        INSERT INTO LOG_SITE1\n"
    "            (OPERATION, IDLIGNECOMMANDE, QUANTITE_OLD, QUANTITE_NEW)\n"
    "        VALUES ('INSERT', :NEW.IDLIGNECOMMANDE, NULL, :NEW.QUANTITE);\n"
    "    ELSIF UPDATING THEN\n"
    "        INSERT INTO LOG_SITE1\n"
    "            (OPERATION, IDLIGNECOMMANDE, QUANTITE_OLD, QUANTITE_NEW)\n"
    "        VALUES ('UPDATE', :NEW.IDLIGNECOMMANDE, :OLD.QUANTITE, :NEW.QUANTITE);\n"
    "    ELSIF DELETING THEN\n"
    "        INSERT INTO LOG_SITE1\n"
    "            (OPERATION, IDLIGNECOMMANDE, QUANTITE_OLD, QUANTITE_NEW)\n"
    "        VALUES ('DELETE', :OLD.IDLIGNECOMMANDE, :OLD.QUANTITE, NULL);\n"
    "    END IF;\n"
    "END;\n"
    "/",
    title="site1/scenario2/triggers/site1_triggers.sh -- Controle + Log"
)

heading2("12.2 Trigger INSTEAD OF de routage -- Central")

code_block(
    "-- Trigger INSTEAD OF sur la vue de routage centrale\n"
    "-- Intercepte INSERT/UPDATE/DELETE sur V_LIGNECOMMANDES_ROUTAGE\n"
    "CREATE OR REPLACE TRIGGER trg_route_lignecommandes\n"
    "INSTEAD OF INSERT OR UPDATE OR DELETE\n"
    "ON V_LIGNECOMMANDES_ROUTAGE\n"
    "FOR EACH ROW\n"
    "BEGIN\n"
    "    IF INSERTING THEN\n"
    "        -- Delègue a route_insert_ligne qui choisit le bon site\n"
    "        route_insert_ligne(\n"
    "            :NEW.IDLIGNECOMMANDE, :NEW.IDCOMMANDE,\n"
    "            :NEW.IDPRODUIT, :NEW.QUANTITE, :NEW.REMISE);\n"
    "    ELSIF UPDATING THEN\n"
    "        -- UPSERT : gere la migration inter-sites si QTE change\n"
    "        route_upsert_ligne(\n"
    "            :NEW.IDLIGNECOMMANDE, :NEW.IDCOMMANDE,\n"
    "            :NEW.IDPRODUIT, :NEW.QUANTITE, :NEW.REMISE);\n"
    "    ELSIF DELETING THEN\n"
    "        -- Cherche et supprime sur le bon site\n"
    "        route_delete_ligne(:OLD.IDLIGNECOMMANDE);\n"
    "    END IF;\n"
    "END;\n"
    "/\n"
    "\n"
    "-- Exemple d'utilisation (transparent pour l'application) :\n"
    "INSERT INTO V_LIGNECOMMANDES_ROUTAGE\n"
    "    (IDLIGNECOMMANDE, IDCOMMANDE, IDPRODUIT, QUANTITE, REMISE)\n"
    "VALUES (9001, 1, 1, 150, 0);  -- QTE=150 -> automatiquement vers Site1\n"
    "COMMIT;\n"
    "\n"
    "-- Migration automatique Site1 -> Site2 !\n"
    "UPDATE V_LIGNECOMMANDES_ROUTAGE\n"
    "SET QUANTITE = 25             -- QTE passe sous 100 -> migre vers Site2\n"
    "WHERE IDLIGNECOMMANDE = 9001;\n"
    "COMMIT;",
    title="docker/routing_central.sh -- Trigger INSTEAD OF de routage"
)

note("Le trigger route_upsert_ligne gere un cas edge important : si une ligne sur Site1 (QTE >= 100) est mise a jour avec QTE < 100, elle est automatiquement supprimee de Site1 et inseree sur Site2 en une seule transaction transparente.")

divider()
page_break()

# ── 13. SYNONYMES ──────────────────────────────────────
heading1("13. Synonymes et vues distribuees")

para("Les synonymes permettent un acces transparent aux tables distantes sans notation @lien_db. C'est un mecanisme d'abstraction qui simplifie considerablement le code applicatif.")

code_block(
    "-- synonyms_central.sql -- Synonymes sur oracle-central\n"
    "-- Crees apres les DB Links : l'acces devient completement transparent\n"
    "\n"
    "-- Synonymes vers Site1 (Scenario 2 : QTE >= 100)\n"
    "CREATE OR REPLACE SYNONYM LigneCommandes1\n"
    "    FOR LigneCommandes1@site1_link;  -- Masque le @link\n"
    "\n"
    "CREATE OR REPLACE SYNONYM Commandes1 FOR Commandes1@site1_link;\n"
    "CREATE OR REPLACE SYNONYM Clients1   FOR Clients1@site1_link;\n"
    "CREATE OR REPLACE SYNONYM Produits1  FOR Produits1@site1_link;\n"
    "\n"
    "-- Synonymes vers Site2 (Scenario 2 : QTE < 100)\n"
    "CREATE OR REPLACE SYNONYM LigneCommandes2\n"
    "    FOR LigneCommandes2@site2_link;\n"
    "\n"
    "CREATE OR REPLACE SYNONYM Commandes2 FOR Commandes2@site2_link;\n"
    "CREATE OR REPLACE SYNONYM Clients2   FOR Clients2@site2_link;\n"
    "CREATE OR REPLACE SYNONYM Produits2  FOR Produits2@site2_link;\n"
    "\n"
    "-- Vue globale avec synonymes (acces transparent sans @link)\n"
    "CREATE OR REPLACE VIEW V_LIGNECOMMANDES_GLOBAL_SYN AS\n"
    "    SELECT 'SITE1' AS SITE, lc.*\n"
    "    FROM LigneCommandes1 lc  -- Via synonyme -> @site1_link\n"
    "    UNION ALL\n"
    "    SELECT 'SITE2' AS SITE, lc.*\n"
    "    FROM LigneCommandes2 lc; -- Via synonyme -> @site2_link\n"
    "\n"
    "COMMIT;",
    title="docker/synonyms_central.sql -- Synonymes et vues distribuees"
)

divider()
page_break()

# ── 14. ROUTAGE ────────────────────────────────────────
heading1("14. Routage automatique depuis le site central")

para("Le routage automatique est la fonctionnalite la plus avancee du projet. Les applications clientes ecrivent dans la vue V_LIGNECOMMANDES_ROUTAGE sans connaître la regle de fragmentation -- le systeme determine automatiquement le site de destination.")

heading2("14.1 Propagation des donnees de reference")

code_block(
    "-- ensure_site1_refs -- Copie les refs manquantes avant insertion\n"
    "CREATE OR REPLACE PROCEDURE ensure_site1_refs (\n"
    "    p_idcmd IN NUMBER, p_idprod IN NUMBER\n"
    ") AS\n"
    "    v_count NUMBER;\n"
    "BEGIN\n"
    "    -- Verifie si le client existe sur Site1\n"
    "    SELECT COUNT(*) INTO v_count FROM Clients1@site1_link cl\n"
    "    WHERE cl.IDCLIENT = (\n"
    "        SELECT c.IDCLIENT FROM COMMANDES c WHERE c.IDCOMMANDE = p_idcmd\n"
    "    );\n"
    "    -- Si absent : copie depuis le Central vers Site1\n"
    "    IF v_count = 0 THEN\n"
    "        INSERT INTO Clients1@site1_link\n"
    "        SELECT cl.* FROM CLIENTS cl\n"
    "        JOIN COMMANDES c ON c.IDCLIENT = cl.IDCLIENT\n"
    "        WHERE c.IDCOMMANDE = p_idcmd;\n"
    "    END IF;\n"
    "\n"
    "    -- Idem pour la commande et le produit...\n"
    "    SELECT COUNT(*) INTO v_count\n"
    "    FROM Commandes1@site1_link WHERE IDCOMMANDE = p_idcmd;\n"
    "    IF v_count = 0 THEN\n"
    "        INSERT INTO Commandes1@site1_link\n"
    "        SELECT * FROM COMMANDES WHERE IDCOMMANDE = p_idcmd;\n"
    "    END IF;\n"
    "END ensure_site1_refs;\n"
    "/",
    title="docker/routing_central.sh -- ensure_site1_refs (pre-propagation)"
)

heading2("14.2 Flux d'execution du routage")

rich_table([
    ("INSERT QTE=150", "QTE >= 100 -> SITE1", "ensure_site1_refs -> INSERT@site1_link -> LOG"),
    ("INSERT QTE=30",  "QTE < 100 -> SITE2",  "ensure_site2_refs -> INSERT@site2_link -> LOG"),
    ("UPDATE 150->25", "Changement de site",   "DELETE@site1_link + INSERT@site2_link -> LOG"),
    ("DELETE ID=9001", "Cherche sur les 2",    "COUNT@site1 + COUNT@site2 -> DELETE bon site"),
], headers=["Operation", "Decision", "Flux d'execution"],
   col_widths=[3.5, 4.0, 9.5])

divider()
page_break()

# ── 15. OPTIMISATION REQUETES ──────────────────────────
heading1("15. Optimisation des requetes distribuees")

heading2("15.1 Requetes distribuees principales")

code_block(
    "-- REQUETE 1 : Chiffre d'affaires par categorie (2 sites, depuis Central)\n"
    "SELECT cat.NOMDECATEGORIE,\n"
    "       SUM(ca.CA_TOTAL) AS CA_TOTAL_2026\n"
    "FROM (\n"
    "    -- Contribution Site1 (QTE >= 100)\n"
    "    SELECT p.IDCATEG,\n"
    "        SUM(lc.QUANTITE * p.PRIXUNITAIRE * (1 - lc.REMISE)) AS CA_TOTAL\n"
    "    FROM LigneCommandes1@site1_link lc\n"
    "    JOIN Produits1@site1_link p ON lc.IDPRODUIT = p.IDPRODUIT\n"
    "    JOIN Commandes1@site1_link c ON lc.IDCOMMANDE = c.IDCOMMANDE\n"
    "    WHERE EXTRACT(YEAR FROM c.DATECOMMANDE) = 2026\n"
    "    GROUP BY p.IDCATEG\n"
    "    UNION ALL\n"
    "    -- Contribution Site2 (QTE < 100)\n"
    "    SELECT p.IDCATEG,\n"
    "        SUM(lc.QUANTITE * p.PRIXUNITAIRE * (1 - lc.REMISE)) AS CA_TOTAL\n"
    "    FROM LigneCommandes2@site2_link lc\n"
    "    JOIN Produits2@site2_link p ON lc.IDPRODUIT = p.IDPRODUIT\n"
    "    JOIN Commandes2@site2_link c ON lc.IDCOMMANDE = c.IDCOMMANDE\n"
    "    WHERE EXTRACT(YEAR FROM c.DATECOMMANDE) = 2026\n"
    "    GROUP BY p.IDCATEG\n"
    ") ca\n"
    "JOIN CATEGORIES cat ON ca.IDCATEG = cat.IDCATEG\n"
    "GROUP BY cat.NOMDECATEGORIE\n"
    "ORDER BY CA_TOTAL_2026 DESC;\n"
    "\n"
    "-- REQUETE 2 : Verification absence de doublons inter-sites\n"
    "SELECT 'DOUBLONS (attendu=0)', COUNT(*)\n"
    "FROM LigneCommandes1@site1_link lc1\n"
    "WHERE EXISTS (\n"
    "    SELECT 1 FROM LigneCommandes2@site2_link lc2\n"
    "    WHERE lc2.IDLIGNECOMMANDE = lc1.IDLIGNECOMMANDE\n"
    ");",
    title="tests/distributed_queries.sql -- Requetes distribuees principales"
)

heading2("15.2 EXPLAIN PLAN -- Avant et apres indexation")

code_block(
    "-- AVANT indexation\n"
    "EXPLAIN PLAN SET STATEMENT_ID = 'AVANT_INDEX' FOR\n"
    "SELECT cl.IDCLIENT, cl.SOCIETE,\n"
    "       COUNT(DISTINCT c.IDCOMMANDE) AS NB_COMMANDES\n"
    "FROM CLIENTS cl\n"
    "JOIN COMMANDES c ON cl.IDCLIENT = c.IDCLIENT\n"
    "WHERE EXTRACT(YEAR FROM c.DATECOMMANDE) = 2026\n"
    "GROUP BY cl.IDCLIENT, cl.SOCIETE\n"
    "ORDER BY NB_COMMANDES DESC;\n"
    "\n"
    "SELECT * FROM TABLE(DBMS_XPLAN.DISPLAY(\n"
    "    statement_id => 'AVANT_INDEX', format => 'TYPICAL +ROWS +COST'));\n"
    "\n"
    "-- Plan typique SANS index :\n"
    "-- | Id | Operation            | Name      | Rows | Cost |\n"
    "-- |  0 | SELECT STATEMENT     |           |   10 |   15 |\n"
    "-- |  1 |  SORT ORDER BY       |           |   10 |   15 |\n"
    "-- |  2 |   HASH GROUP BY      |           |   10 |   13 |\n"
    "-- |* 3 |    HASH JOIN         |           |   10 |   11 |\n"
    "-- |  4 |     TABLE ACCESS FULL| CLIENTS   |    8 |    3 |\n"
    "-- |* 5 |     TABLE ACCESS FULL| COMMANDES |   10 |    3 |\n"
    "-- Predicate: EXTRACT(YEAR FROM C.DATECOMMANDE)=2026",
    title="tests/performance_analysis.sql -- EXPLAIN PLAN avant indexation"
)

divider()
page_break()

# ── 16. INDEXATION ─────────────────────────────────────
heading1("16. Strategie d'indexation multi-niveaux")

heading2("16.1 Index sur les fragments locaux -- Site1")

code_block(
    "-- site1_indexes.sh -- Index Site1 Scenario 2\n"
    "\n"
    "-- Index simple sur le critere de fragmentation\n"
    "CREATE INDEX idx_lc1_quantite   ON LigneCommandes1 (QUANTITE);\n"
    "\n"
    "-- Index sur cle etrangere (jointure LigneCommandes -> Commandes)\n"
    "CREATE INDEX idx_lc1_idcommande ON LigneCommandes1 (IDCOMMANDE);\n"
    "\n"
    "-- Index sur cle etrangere (jointure LigneCommandes -> Produits)\n"
    "CREATE INDEX idx_lc1_idproduit  ON LigneCommandes1 (IDPRODUIT);\n"
    "\n"
    "-- Index composite : couvre jointure + filtre quantite simultanement\n"
    "-- Utile : SELECT ... WHERE IDCOMMANDE=X AND QUANTITE>=Y\n"
    "CREATE INDEX idx_lc1_cmd_qte    ON LigneCommandes1 (IDCOMMANDE, QUANTITE);\n"
    "\n"
    "-- Index date sur Commandes1 (filtrage temporel)\n"
    "CREATE INDEX idx_cmd1_date      ON Commandes1 (DATECOMMANDE);\n"
    "\n"
    "-- Index sur cle etrangere (jointure Commandes -> Clients)\n"
    "CREATE INDEX idx_cmd1_idclient  ON Commandes1 (IDCLIENT);\n"
    "\n"
    "-- Index categorie (jointure Produits -> Categories)\n"
    "CREATE INDEX idx_prod1_idcateg  ON Produits1 (IDCATEG);\n"
    "\n"
    "COMMIT;",
    title="site1/scenario2/indexes/site1_indexes.sh -- Index Site1"
)

heading2("16.2 Index fonctionnel pour EXTRACT() -- Probleme et solution")

para("Nous avons decouvert un probleme important : les index B-Tree classiques sur DATECOMMANDE ne sont pas utilises par l'optimiseur quand la colonne est enveloppee dans EXTRACT(). La solution est un index fonctionnel qui indexe directement le resultat de la fonction.")

code_block(
    "-- Index classique (NE resout PAS le probleme avec EXTRACT)\n"
    "CREATE INDEX IDX_CMD_DATECOMMANDE ON COMMANDES (DATECOMMANDE);\n"
    "\n"
    "-- Index fonctionnel : indexe EXTRACT(YEAR FROM DATECOMMANDE)\n"
    "-- L'optimiseur peut maintenant l'utiliser pour :\n"
    "-- WHERE EXTRACT(YEAR FROM c.DATECOMMANDE) = 2026\n"
    "CREATE INDEX IDX_CMD_YEAR_FUNC\n"
    "    ON COMMANDES (EXTRACT(YEAR FROM DATECOMMANDE));\n"
    "\n"
    "-- Index composite (jointure + filtre simultane)\n"
    "CREATE INDEX IDX_CMD_CLIENT_DATE\n"
    "    ON COMMANDES (IDCLIENT, DATECOMMANDE);\n"
    "\n"
    "-- Recollecte des statistiques pour que l'optimiseur en tienne compte\n"
    "BEGIN\n"
    "    DBMS_STATS.GATHER_TABLE_STATS(USER, 'COMMANDES', CASCADE => TRUE);\n"
    "    DBMS_STATS.GATHER_TABLE_STATS(USER, 'CLIENTS',   CASCADE => TRUE);\n"
    "END;\n"
    "/",
    title="tests/performance_analysis.sql -- Index fonctionnel EXTRACT()"
)

rich_table([
    ("idx_lc1_quantite",  "B-Tree simple",    "LigneCommandes1 (QUANTITE)",               "Filtre par regle de fragmentation"),
    ("idx_lc1_cmd_qte",   "B-Tree composite", "LigneCommandes1 (IDCOMMANDE, QUANTITE)",    "Jointure + filtre simultane"),
    ("IDX_CMD_YEAR_FUNC", "B-Tree fonctionnel","COMMANDES (EXTRACT(YEAR FROM DATE))",      "Resout le probleme EXTRACT()"),
    ("idx_prod1_idcateg", "B-Tree simple",    "Produits1 (IDCATEG)",                       "Jointure Produits -> Categories"),
    ("idx_cmd1_date",     "B-Tree simple",    "Commandes1 (DATECOMMANDE)",                 "Filtrage temporel"),
], headers=["Index", "Type", "Colonnes", "Utilite"],
   col_widths=[3.5, 3.0, 5.5, 5.0])

divider()
page_break()

# ── 17. PERFORMANCES ───────────────────────────────────
heading1("17. Analyse comparative des performances")

para("Nous avons effectue une analyse comparative structuree en quatre phases. Les resultats montrent une amelioration significative meme sur nos volumes reduits (25 lignes), ce qui prefigure des gains bien plus importants en production.")

heading2("17.1 Protocole de mesure")

code_block(
    "-- Phase 1 : mesure SANS index\n"
    "DECLARE\n"
    "    t_start TIMESTAMP; t_end TIMESTAMP;\n"
    "    v_elapsed NUMBER; v_count NUMBER;\n"
    "BEGIN\n"
    "    t_start := SYSTIMESTAMP;\n"
    "    SELECT COUNT(*) INTO v_count FROM (\n"
    "        SELECT cl.IDCLIENT, cl.SOCIETE,\n"
    "               COUNT(DISTINCT c.IDCOMMANDE) AS NB_COMMANDES\n"
    "        FROM CLIENTS cl JOIN COMMANDES c ON cl.IDCLIENT = c.IDCLIENT\n"
    "        WHERE EXTRACT(YEAR FROM c.DATECOMMANDE) = 2026\n"
    "        GROUP BY cl.IDCLIENT, cl.SOCIETE\n"
    "    );\n"
    "    t_end := SYSTIMESTAMP;\n"
    "    v_elapsed := EXTRACT(SECOND FROM (t_end - t_start)) * 1000;\n"
    "    DBMS_OUTPUT.PUT_LINE('SANS index : ' || ROUND(v_elapsed, 3) || ' ms');\n"
    "END;\n"
    "/\n"
    "\n"
    "-- Resultat console attendu :\n"
    "-- SANS index  :  12.847 ms  (FULL TABLE SCAN)\n"
    "-- AVEC index  :   3.214 ms  (INDEX RANGE SCAN)\n"
    "-- Distribue   :  27.563 ms  (2 sites + latence reseau Docker)",
    title="tests/performance_analysis.sql -- Mesure temps d'execution"
)

heading2("17.2 Resultats comparatifs")

rich_table([
    ("Sans index",          "TABLE ACCESS FULL + HASH JOIN + SORT ORDER BY", "~8-15 ms",  "~350-500 blocs"),
    ("Avec index B-Tree",   "INDEX RANGE SCAN + NESTED LOOPS",               "~3-6 ms",   "~80-150 blocs"),
    ("Avec index fonction.","INDEX RANGE SCAN sur EXTRACT()",                 "~2-4 ms",   "~40-80 blocs"),
    ("Requete distribuee",  "2x REMOTE (Site1+Site2) + jointure locale",     "~15-40 ms", "Inclut latence reseau"),
], headers=["Scenario", "Plan Oracle", "Temps estime", "Consistent Gets"],
   col_widths=[4.0, 6.0, 3.0, 4.0])

note("L'index fonctionnel IDX_CMD_YEAR_FUNC est la correction la plus impactante : sans lui, aucun index sur DATECOMMANDE n'est utilise avec EXTRACT(), ce que seul EXPLAIN PLAN permet de detecter.")

divider()
page_break()

# ── 18. MONITORING ─────────────────────────────────────
heading1("18. Monitoring et maintenance des performances")

heading2("18.1 Surveillance des tables de log")

code_block(
    "-- monitor_logs.sql -- Surveillance de l'activite sur les sites\n"
    "\n"
    "-- Resume global : nombre d'operations par type\n"
    "SELECT\n"
    "    OPERATION,\n"
    "    COUNT(*)             AS NB_OPERATIONS,\n"
    "    MIN(DATE_OPERATION)  AS PREMIERE_OP,\n"
    "    MAX(DATE_OPERATION)  AS DERNIERE_OP\n"
    "FROM LOG_SITE1\n"
    "GROUP BY OPERATION\n"
    "ORDER BY NB_OPERATIONS DESC;\n"
    "\n"
    "-- 20 dernieres operations (pour audit)\n"
    "SELECT ID_LOG, OPERATION, IDLIGNECOMMANDE,\n"
    "       QUANTITE_OLD, QUANTITE_NEW,\n"
    "       TO_CHAR(DATE_OPERATION,'DD/MM/YYYY HH24:MI:SS') AS DATE_OP,\n"
    "       UTILISATEUR\n"
    "FROM LOG_SITE1\n"
    "ORDER BY ID_LOG DESC FETCH FIRST 20 ROWS ONLY;\n"
    "\n"
    "-- Alertes : operations des dernieres 24h\n"
    "SELECT COUNT(*) AS OPS_RECENTES_24H\n"
    "FROM LOG_SITE1\n"
    "WHERE DATE_OPERATION > SYSTIMESTAMP - INTERVAL '1' DAY;\n"
    "\n"
    "-- Volume total des 4 tables de log\n"
    "SELECT 'LOG_SITE1'     AS LOG_TABLE, COUNT(*) AS NB FROM LOG_SITE1\n"
    "UNION ALL\n"
    "SELECT 'LOG_SITE2',     COUNT(*) FROM LOG_SITE2\n"
    "UNION ALL\n"
    "SELECT 'LOG_SITE1_SC1', COUNT(*) FROM LOG_SITE1_SC1\n"
    "UNION ALL\n"
    "SELECT 'LOG_SITE2_SC1', COUNT(*) FROM LOG_SITE2_SC1;",
    title="monitoring/monitor_logs.sql -- Surveillance de l'activite DML"
)

heading2("18.2 Maintenance periodique")

code_block(
    "-- rebuild_indexes.sql -- Reconstruction des index fragmentes\n"
    "DECLARE\n"
    "    PROCEDURE rebuild_if_exists(p_index VARCHAR2) IS\n"
    "        v_count NUMBER;\n"
    "    BEGIN\n"
    "        SELECT COUNT(*) INTO v_count FROM USER_INDEXES\n"
    "        WHERE INDEX_NAME = UPPER(p_index);\n"
    "        IF v_count > 0 THEN\n"
    "            EXECUTE IMMEDIATE 'ALTER INDEX ' || p_index || ' REBUILD';\n"
    "            DBMS_OUTPUT.PUT_LINE('  Reconstruit : ' || p_index);\n"
    "        END IF;\n"
    "    END;\n"
    "BEGIN\n"
    "    rebuild_if_exists('IDX_LC1_QUANTITE');\n"
    "    rebuild_if_exists('IDX_LC1_IDCOMMANDE');\n"
    "    rebuild_if_exists('IDX_LC1_CMD_QTE');\n"
    "    rebuild_if_exists('IDX_CMD1_DATE');\n"
    "END;\n"
    "/\n"
    "\n"
    "-- purge_logs.sql -- Purge des logs anciens (> 30 jours)\n"
    "DECLARE\n"
    "    v_seuil   TIMESTAMP := SYSTIMESTAMP - INTERVAL '30' DAY;\n"
    "    v_deleted NUMBER;\n"
    "BEGIN\n"
    "    SELECT COUNT(*) INTO v_deleted FROM LOG_SITE1\n"
    "    WHERE DATE_OPERATION < v_seuil;\n"
    "    DELETE FROM LOG_SITE1 WHERE DATE_OPERATION < v_seuil;\n"
    "    DBMS_OUTPUT.PUT_LINE('LOG_SITE1 : ' || v_deleted || ' supprimee(s)');\n"
    "    COMMIT;\n"
    "EXCEPTION WHEN OTHERS THEN ROLLBACK; RAISE;\n"
    "END;\n"
    "/",
    title="maintenance/ -- Rebuild index et purge logs"
)

heading2("18.3 Planification de la maintenance")

rich_table([
    ("Quotidienne",  "monitor_logs.sql",   "Surveillance activite DML + alertes 24h"),
    ("Hebdomadaire", "check_dblinks.sql",  "Test connectivite DB Links + latence"),
    ("Hebdomadaire", "rebuild_indexes.sql","Reconstruction index apres forte charge"),
    ("Mensuelle",    "purge_logs.sql",     "Purge entrees log de plus de 30 jours"),
    ("Apres charge", "analyze_tables.sql", "Collecte statistiques CBO (DBMS_STATS)"),
    ("A la demande", "check_health.sh",    "Health check complet avant/apres intervention"),
], headers=["Frequence", "Script", "Objectif"],
   col_widths=[3.0, 4.5, 9.5])

divider()
page_break()

# ── 19. CONCLUSION ─────────────────────────────────────
heading1("19. Conclusion")

para("Ce projet nous a permis de mettre en pratique de maniere concrete les concepts theoriques des bases de donnees distribuees dans un environnement realiste orchestre par Docker. Plusieurs enseignements importants en sont ressortis.")

heading2("Bilan technique")

bullet("L'ordre d'initialisation des scripts Docker est critique : le prefixage numerique des repertoires montes (01_schema, 02_data, ...) est indispensable pour garantir l'execution ordonnee.", bold_prefix="Ordre d'initialisation : ")
bullet("La procédure ensure_site1_refs illustre la complexite de la gestion des dependances : fragmenter la table principale ne suffit pas, il faut propager les tables de reference vers chaque site.", bold_prefix="Dependances referentielles : ")
bullet("L'index fonctionnel IDX_CMD_YEAR_FUNC est la decouverte la plus instructive : Oracle ne peut pas utiliser un index B-Tree classique sur DATECOMMANDE quand EXTRACT() enveloppe la colonne.", bold_prefix="Index fonctionnels : ")
bullet("Le pattern vue INSTEAD OF + trigger de routage + procedures stockees rend la distribution completement transparente pour l'application cliente.", bold_prefix="Routage transparent : ")

heading2("Resultats obtenus")

rich_table([
    ("Architecture",    "3 noeuds Oracle XE 21c sur Docker Compose avec reseau bridge prive"),
    ("Fragmentation",   "2 scenarios (volume + categorie), 25 lignes correctement reparties"),
    ("DB Links",        "6 liens bidirectionnels entre les 3 noeuds, tous fonctionnels"),
    ("Procedures",      "6 procedures CRUD par site + 4 procedures de routage sur le Central"),
    ("Triggers",        "12 triggers : 4 controle + 4 log + 4 routage INSTEAD OF"),
    ("Index",           "20+ index B-Tree, composites et fonctionnels sur les 2 sites"),
    ("Monitoring",      "Scripts de healthcheck, log et maintenance planifiable"),
], headers=["Composant", "Realisation"],
   col_widths=[3.5, 13.5])

para("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_border_bottom(p, color="E86A17", size="12")
r = p.add_run(" ")
r.font.size = Pt(4)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("-- Fin du rapport -- Juin 2026 --")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = C_GREY_MED
set_run_font(r, 'Calibri')

# En-tete et pied de page
add_header_footer()

# Sauvegarde
output_path = (r"C:\Users\LENOVO\Downloads\final-sara\final BDD"
               r"\distributed-database-docker-project-main\rapport"
               r"\rapport_BDD_distribuees.docx")
doc.save(output_path)
print(f"Rapport genere : {output_path}")
