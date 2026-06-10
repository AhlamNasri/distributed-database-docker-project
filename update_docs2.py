#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Ajoute la section sur la correction de la redondance LIGNECOMMANDES
dans les deux documents Word.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLEU_FONCE = RGBColor(0x0D, 0x1B, 0x3E)
ORANGE     = RGBColor(0xE8, 0x6B, 0x1A)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_para(doc, text='', bold=False, italic=False, size=11,
             color=None, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p

def add_bullet(doc, text, bold_prefix=None, size=10.5, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.8)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(size)
        r2 = p.add_run(text); r2.font.size = Pt(size)
    else:
        r = p.add_run(text); r.font.size = Pt(size)
    return p

def add_code(doc, code, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    r = p.add_run(code)
    r.font.name = 'Courier New'; r.font.size = Pt(size)
    return p

def add_divider(doc, color='E86B1A'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)

def build_redondance_section(doc, section_number):
    """Ajoute la section sur la correction de redondance."""

    # ── Titre ────────────────────────────────────────────────────────────────
    h = doc.add_heading(
        f'{section_number}. CORRECTION DE LA REDONDANCE — TRUNCATE LIGNECOMMANDES',
        level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after  = Pt(8)
    for run in h.runs:
        run.font.color.rgb = BLEU_FONCE
    add_divider(doc)

    # ── Identification ───────────────────────────────────────────────────────
    add_para(doc, f'{section_number}.1 Identification du problème',
             bold=True, color=ORANGE, size=12, space_after=4)

    add_para(doc,
        "Lors de l'analyse de l'ordre d'initialisation du cluster, un second problème "
        "de redondance a été identifié, distinct du problème de doublon corrigé "
        "précédemment. Ce problème est structurel et concerne les données initiales "
        "insérées au démarrage.")

    add_para(doc,
        "Le script eshop_data.sh est monté sur les trois conteneurs (Site1, Site2 et "
        "Central). À l'étape 02, les 25 lignes de commande sont donc insérées "
        "simultanément dans la table LIGNECOMMANDES de chaque nœud. À l'étape 05, "
        "chaque site crée ses tables fragmentées via CREATE TABLE AS SELECT, "
        "extrayant la portion qui lui revient. Cependant, la table source "
        "LIGNECOMMANDES n'est jamais vidée après cette opération.")

    add_para(doc, "Résultat avant correction :", bold=True, size=10, space_after=3)

    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_txt in enumerate(["Nœud", "Table", "Lignes"]):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, '1A3F7A')
        r = cell.paragraphs[0].add_run(h_txt)
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows_data = [
        ["Site1",   "LIGNECOMMANDES (redondante)",  "25 lignes — inutiles"],
        ["Site1",   "LigneCommandes1 (fragment)",   "12 lignes — actives"],
        ["Site2",   "LIGNECOMMANDES (redondante)",  "25 lignes — inutiles"],
    ]
    bgs = ['FFF0F0', 'F0FFF4', 'FFF0F0']
    for ri, (row_data, bg) in enumerate(zip(rows_data, bgs)):
        for ci, val in enumerate(row_data):
            cell = tbl.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, '')

    add_para(doc,
        "Les 25 lignes dans LIGNECOMMANDES coexistaient silencieusement avec les "
        "tables fragmentées sur chaque nœud. Bien que non utilisées par les "
        "procédures et triggers (qui travaillent uniquement sur LigneCommandes1 "
        "et LigneCommandes2), elles représentaient une incohérence structurelle : "
        "les mêmes données existaient à deux endroits différents sur chaque site.")

    # ── Solution ─────────────────────────────────────────────────────────────
    add_para(doc, f'{section_number}.2 Solution appliquée',
             bold=True, color=ORANGE, size=12, space_after=4)

    add_para(doc,
        "La solution consiste à ajouter une instruction TRUNCATE TABLE LIGNECOMMANDES "
        "à la fin de chaque script de fragmentation, immédiatement après la création "
        "des tables fragmentées. Ainsi, dès que les fragments sont créés et peuplés, "
        "la table source est vidée. Trois fichiers ont été modifiés :")

    add_para(doc, "Fichier 1 — site1/scenario2/fragments/site1_fragments.sh :",
             bold=True, size=10, space_after=3)
    add_code(doc,
        "-- Après CREATE TABLE LigneCommandes1 AS SELECT ... WHERE QUANTITE >= 100\n"
        "ALTER TABLE LigneCommandes1 ADD CONSTRAINT chk_sc2_site1_qte\n"
        "    CHECK (QUANTITE >= 100);\n\n"
        "-- AJOUT : vider la table source devenue redondante\n"
        "TRUNCATE TABLE LIGNECOMMANDES;\n\n"
        "COMMIT;")

    add_para(doc, "Fichier 2 — site2/scenario2/fragments/site2_fragments.sh :",
             bold=True, size=10, space_after=3)
    add_code(doc,
        "-- Après CREATE TABLE LigneCommandes2 AS SELECT ... WHERE QUANTITE < 100\n"
        "ALTER TABLE LigneCommandes2 ADD CONSTRAINT chk_sc2_site2_qte\n"
        "    CHECK (QUANTITE < 100);\n\n"
        "-- AJOUT : vider la table source devenue redondante\n"
        "TRUNCATE TABLE LIGNECOMMANDES;\n\n"
        "COMMIT;")

    add_para(doc, "Fichier 3 — docker/routing_central.sh :",
             bold=True, size=10, space_after=3)
    add_code(doc,
        "-- En fin de script, après la création de toute la couche de routage\n"
        "-- AJOUT : vider LIGNECOMMANDES sur le Central\n"
        "-- Le routage via V_LIGNECOMMANDES_ROUTAGE est désormais le seul point d'entrée\n"
        "TRUNCATE TABLE LIGNECOMMANDES;\n\n"
        "COMMIT;")

    add_para(doc,
        "Le TRUNCATE est placé après la création des fragments (et non avant) car "
        "le CREATE TABLE AS SELECT a besoin de lire LIGNECOMMANDES pour extraire "
        "les données. Sur le Central, le TRUNCATE est placé en fin de "
        "routing_central.sh, dernier script d'initialisation, garantissant que "
        "toute la couche de routage est opérationnelle avant de vider la table.")

    # ── Résultats ─────────────────────────────────────────────────────────────
    add_para(doc, f'{section_number}.3 Résultats de validation',
             bold=True, color=ORANGE, size=12, space_after=4)

    add_para(doc,
        "Après redéploiement complet du cluster (docker-compose down -v && "
        "docker-compose up -d), les vérifications suivantes ont été effectuées "
        "depuis oracle-central :")

    add_code(doc,
        "=== LIGNECOMMANDES Central (doit etre vide) ===\n"
        "LIGNECOMMANDES_CENTRAL\n"
        "----------------------\n"
        "                     0        <-- vide, zero redondance\n"
        "\n"
        "=== Fragment Site1 via DB Link (doit avoir 12 lignes) ===\n"
        "LIGNECOMMANDES1_SITE1\n"
        "---------------------\n"
        "                   12        <-- correct\n"
        "\n"
        "=== Fragment Site2 via DB Link (doit avoir 13 lignes) ===\n"
        "LIGNECOMMANDES2_SITE2\n"
        "---------------------\n"
        "                   13        <-- correct\n"
        "\n"
        "=== Vue globale (doit avoir 25 lignes au total) ===\n"
        "SITE   NB_LIGNES\n"
        "----- ----------\n"
        "SITE1         12\n"
        "SITE2         13             <-- total = 25, zero perte de donnees")

    # ── Tableau comparatif ───────────────────────────────────────────────────
    add_para(doc, '')
    add_para(doc, "Tableau comparatif avant / après correction :",
             bold=True, size=10, space_after=3)

    tbl2 = doc.add_table(rows=5, cols=3)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_txt in enumerate(["Table / Nœud", "Avant correction", "Après correction"]):
        cell = tbl2.cell(0, i)
        set_cell_bg(cell, '0D4E2E')
        r = cell.paragraphs[0].add_run(h_txt)
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows2 = [
        ["LIGNECOMMANDES — Central",  "25 lignes (redondantes)", "0 ligne ✅"],
        ["LIGNECOMMANDES — Site1",    "25 lignes (redondantes)", "0 ligne ✅"],
        ["LIGNECOMMANDES — Site2",    "25 lignes (redondantes)", "0 ligne ✅"],
        ["LigneCommandes1 — Site1",   "12 lignes",               "12 lignes ✅"],
    ]
    bgs2 = ['FFF0F0', 'FFF0F0', 'FFF0F0', 'F0FFF4']
    for ri, (row_data, bg) in enumerate(zip(rows2, bgs2)):
        for ci, val in enumerate(row_data):
            cell = tbl2.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10)
            if ci == 0:
                r.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, '')

    # ── Impact ───────────────────────────────────────────────────────────────
    add_para(doc, f'{section_number}.4 Impact sur l\'architecture',
             bold=True, color=ORANGE, size=12, space_after=4)

    add_bullet(doc,
        "Chaque ligne de commande n'existe plus qu'à un seul endroit physique "
        "(Site1 OU Site2), conformément au principe de la fragmentation horizontale.",
        bold_prefix="Unicité des données : ")
    add_bullet(doc,
        "La table LIGNECOMMANDES reste présente sur chaque nœud (structure "
        "conservée) mais vide — elle sert de point d'entrée protégé par le "
        "trigger BEFORE INSERT.",
        bold_prefix="Structure conservée : ")
    add_bullet(doc,
        "Ce fix complète le fix précédent (trigger BEFORE INSERT) : "
        "ensemble, ils garantissent qu'aucune donnée ne peut exister en double, "
        "ni au démarrage ni en cours d'utilisation.",
        bold_prefix="Complémentarité : ")

    add_para(doc, '')


# ═══════════════════════════════════════════════════════════════════════════════
# 1. Mise à jour du rapport
# ═══════════════════════════════════════════════════════════════════════════════

rapport_path = (r"C:\Users\LENOVO\Downloads\project-sara-aham"
                r"\distributed-database-docker-project-main"
                r"\rapport\rapport_BDD_distribuees (1).docx")

doc_rapport = Document(rapport_path)

# Ajouter entrée dans la TDM après la section 20
for para in doc_rapport.paragraphs:
    if 'Correction du probl' in para.text and 'doublon' in para.text.lower():
        new_toc = doc_rapport.add_paragraph()
        para._p.addnext(new_toc._p)
        r = new_toc.add_run(
            "  21. Correction de la redondance — TRUNCATE LIGNECOMMANDES")
        r.font.size = Pt(11)
        break

doc_rapport.add_page_break()
build_redondance_section(doc_rapport, 21)
doc_rapport.save(rapport_path)
print("[OK] Rapport mis à jour : rapport_BDD_distribuees (1).docx")


# ═══════════════════════════════════════════════════════════════════════════════
# 2. Mise à jour du document d'explication
# ═══════════════════════════════════════════════════════════════════════════════

expl_path = (r"C:\Users\LENOVO\Downloads\project-sara-aham"
             r"\distributed-database-docker-project-main"
             r"\rapport\Explication_PPT_BDD_AhlamNasri_SaraHimmad.docx")

doc_expl = Document(expl_path)

# Ajouter entrée dans le plan après section 11
for para in doc_expl.paragraphs:
    if 'Correction du Probl' in para.text and 'Doublon' in para.text:
        new_plan = doc_expl.add_paragraph(style='List Bullet')
        para._p.addnext(new_plan._p)
        r = new_plan.add_run(
            "12 · Correction de la Redondance — TRUNCATE LIGNECOMMANDES")
        r.font.size = Pt(10.5)
        break

doc_expl.add_page_break()
build_redondance_section(doc_expl, 12)
doc_expl.save(expl_path)
print("[OK] Explication mise à jour : Explication_PPT_BDD_AhlamNasri_SaraHimmad.docx")

print("\nMise à jour terminée.")
