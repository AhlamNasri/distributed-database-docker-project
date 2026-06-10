#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Ajoute la section sur la correction de la faille de doublons sur Site1/Site2
(triggers BEFORE INSERT sur LIGNECOMMANDES) dans les deux documents Word.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLEU_FONCE = RGBColor(0x0D, 0x1B, 0x3E)
ORANGE     = RGBColor(0xE8, 0x6B, 0x1A)
VERT       = RGBColor(0x0D, 0x4E, 0x2E)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_para(doc, text='', bold=False, italic=False, size=11,
             color=None, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p

def add_bullet(doc, text, bold_prefix=None, size=10.5, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.8)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(size)
        r2 = p.add_run(text); r2.font.size = Pt(size)
    else:
        r = p.add_run(text); r.font.size = Pt(size)
    return p

def add_code(doc, code, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    r = p.add_run(code)
    r.font.name = 'Courier New'; r.font.size = Pt(size)
    return p

def add_divider(doc, color='E86B1A'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)

def build_section(doc, section_number):
    # ── Titre ────────────────────────────────────────────────────────────────
    h = doc.add_heading(
        f'{section_number}. CORRECTION DE LA FAILLE DE DOUBLONS SUR SITE1 ET SITE2',
        level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after  = Pt(8)
    for run in h.runs:
        run.font.color.rgb = BLEU_FONCE
    add_divider(doc)

    # ── Contexte ─────────────────────────────────────────────────────────────
    add_para(doc, f'{section_number}.1 Contexte — couverture incomplète du Fix #1',
             bold=True, color=ORANGE, size=12, space_after=4)

    add_para(doc,
        "Le Fix #1 avait corrigé le problème de doublon sur le Central en remplaçant "
        "le trigger AFTER par un trigger BEFORE qui bloque tout DML direct sur "
        "LIGNECOMMANDES (ORA-20300). Cependant, cette protection ne couvrait que "
        "le noeud Central. Sur Site1 et Site2, la table LIGNECOMMANDES était certes "
        "vide (grâce au TRUNCATE du Fix #2), mais aucun trigger ne bloquait "
        "un éventuel INSERT direct. Ce scénario restait possible :")

    add_code(doc,
        "-- Depuis Site1, un INSERT direct passait sans erreur\n"
        "INSERT INTO LIGNECOMMANDES (IDLIGNECOMMANDE, IDCOMMANDE, IDPRODUIT, QUANTITE, REMISE)\n"
        "VALUES (9998, 1, 1, 150, 0);\n\n"
        "-- Résultat AVANT ce fix : 1 row created (donnée hors fragment, sans routage)")

    add_para(doc,
        "La ligne s'inscrivait dans LIGNECOMMANDES sur Site1 sans être routée "
        "vers LigneCommandes1 ni LigneCommandes2, et sans apparaître dans la vue "
        "V_LIGNECOMMANDES_GLOBAL_SYN. C'est la même cause racine que le Fix #1, "
        "mais sur les nœuds locaux.")

    # ── Diagnostic ───────────────────────────────────────────────────────────
    add_para(doc, f'{section_number}.2 Diagnostic — état avant correction',
             bold=True, color=ORANGE, size=12, space_after=4)

    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_txt in enumerate(["Noeud", "Table LIGNECOMMANDES", "Protection"]):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, '1A3F7A')
        r = cell.paragraphs[0].add_run(h_txt)
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows_data = [
        ["Central", "LIGNECOMMANDES vide",   "BEFORE trigger ORA-20300 (Fix #1) ✅"],
        ["Site1",   "LIGNECOMMANDES vide",   "Aucune protection ❌"],
        ["Site2",   "LIGNECOMMANDES vide",   "Aucune protection ❌"],
    ]
    bgs = ['E8F5E9', 'FFEBEE', 'FFEBEE']
    for ri, (row_data, bg) in enumerate(zip(rows_data, bgs)):
        for ci, val in enumerate(row_data):
            cell = tbl.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10)
            if ci == 0:
                r.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, '')

    # ── Solution ─────────────────────────────────────────────────────────────
    add_para(doc, f'{section_number}.3 Solution — triggers BEFORE sur Site1 et Site2',
             bold=True, color=ORANGE, size=12, space_after=4)

    add_para(doc,
        "La correction consiste à ajouter un trigger BEFORE INSERT OR UPDATE OR DELETE "
        "sur LIGNECOMMANDES dans les fichiers de triggers de chaque site, sur le même "
        "modèle que le trigger du Central. Deux fichiers ont été modifiés :")

    add_para(doc, "Fichier 1 — site1/scenario2/triggers/site1_triggers.sh :",
             bold=True, size=10, space_after=3)
    add_code(doc,
        "-- Ajouté à la fin du bloc SQL, avant COMMIT/EXIT\n"
        "CREATE OR REPLACE TRIGGER trg_block_lignecommandes_site1\n"
        "BEFORE INSERT OR UPDATE OR DELETE ON LIGNECOMMANDES\n"
        "FOR EACH ROW\n"
        "BEGIN\n"
        "    RAISE_APPLICATION_ERROR(-20301,\n"
        "        'DML direct sur LIGNECOMMANDES interdit sur Site1. '\n"
        "        || 'Utilisez la vue V_LIGNECOMMANDES_ROUTAGE sur le Central.');\n"
        "END;\n"
        "/")

    add_para(doc, "Fichier 2 — site2/scenario2/triggers/site2_triggers.sh :",
             bold=True, size=10, space_after=3)
    add_code(doc,
        "-- Ajouté à la fin du bloc SQL, avant COMMIT/EXIT\n"
        "CREATE OR REPLACE TRIGGER trg_block_lignecommandes_site2\n"
        "BEFORE INSERT OR UPDATE OR DELETE ON LIGNECOMMANDES\n"
        "FOR EACH ROW\n"
        "BEGIN\n"
        "    RAISE_APPLICATION_ERROR(-20302,\n"
        "        'DML direct sur LIGNECOMMANDES interdit sur Site2. '\n"
        "        || 'Utilisez la vue V_LIGNECOMMANDES_ROUTAGE sur le Central.');\n"
        "END;\n"
        "/")

    # ── Résultats ─────────────────────────────────────────────────────────────
    add_para(doc, f'{section_number}.4 Résultats de validation',
             bold=True, color=ORANGE, size=12, space_after=4)

    add_para(doc,
        "Après redéploiement du cluster, le test de blocage a été exécuté sur Site1 :")

    add_code(doc,
        "=== Test INSERT direct sur Site1 (doit etre bloque ORA-20301) ===\n"
        "\n"
        "INSERT INTO LIGNECOMMANDES (IDLIGNECOMMANDE, IDCOMMANDE,\n"
        "                            IDPRODUIT, QUANTITE, REMISE)\n"
        "VALUES (9998, 1, 1, 150, 0);\n"
        "            *\n"
        "ERROR at line 1:\n"
        "ORA-20301: DML direct sur LIGNECOMMANDES interdit sur Site1.\n"
        "           Utilisez la vue V_LIGNECOMMANDES_ROUTAGE sur le Central.\n"
        "ORA-06512: at \"ESHOP1.TRG_BLOCK_LIGNECOMMANDES_SITE1\", line 2\n"
        "ORA-04088: error during execution of trigger\n"
        "           'ESHOP1.TRG_BLOCK_LIGNECOMMANDES_SITE1'")

    add_para(doc, '')

    # ── Tableau final ─────────────────────────────────────────────────────────
    add_para(doc, "Tableau de protection final — les 3 noeuds :",
             bold=True, size=10, space_after=3)

    tbl2 = doc.add_table(rows=5, cols=4)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_txt in enumerate(["Noeud", "Table protegee", "Trigger", "Code erreur"]):
        cell = tbl2.cell(0, i)
        set_cell_bg(cell, '0D4E2E')
        r = cell.paragraphs[0].add_run(h_txt)
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows2 = [
        ["Central", "LIGNECOMMANDES",              "trg_route_lignecommandes_table",   "ORA-20300"],
        ["Site1",   "LIGNECOMMANDES",              "trg_block_lignecommandes_site1",   "ORA-20301"],
        ["Site2",   "LIGNECOMMANDES",              "trg_block_lignecommandes_site2",   "ORA-20302"],
        ["Central", "V_LIGNECOMMANDES_ROUTAGE",    "INSTEAD OF (routage automatique)", "—"],
    ]
    bgs2 = ['FFF8E1', 'E8F5E9', 'E8F5E9', 'E3F2FD']
    for ri, (row_data, bg) in enumerate(zip(rows2, bgs2)):
        for ci, val in enumerate(row_data):
            cell = tbl2.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(9.5)
            if ci == 0:
                r.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, '')

    # ── Impact ───────────────────────────────────────────────────────────────
    add_para(doc, f'{section_number}.5 Impact et complémentarité des trois fixes',
             bold=True, color=ORANGE, size=12, space_after=4)

    add_bullet(doc,
        "Bloque les inserts directs sur LIGNECOMMANDES du Central (ORA-20300) — "
        "oblige à passer par la vue de routage.",
        bold_prefix="Fix #1 — Trigger Central : ")
    add_bullet(doc,
        "Vide LIGNECOMMANDES sur les 3 noeuds après la création des fragments — "
        "élimine la redondance initiale des 25 lignes.",
        bold_prefix="Fix #2 — TRUNCATE : ")
    add_bullet(doc,
        "Bloque les inserts directs sur LIGNECOMMANDES de Site1/Site2 (ORA-20301/20302) — "
        "ferme la faille résiduelle sur les noeuds locaux.",
        bold_prefix="Fix #3 — Triggers Site1/Site2 : ")

    add_para(doc,
        "Ensemble, ces trois corrections garantissent qu'aucune ligne de commande "
        "ne peut exister en dehors des fragments LigneCommandes1 et LigneCommandes2, "
        "ni au démarrage ni en cours d'utilisation, quel que soit le noeud "
        "depuis lequel l'utilisateur tente d'écrire.")

    add_para(doc, '')


# ═══════════════════════════════════════════════════════════════════════════════
# 1. Rapport
# ═══════════════════════════════════════════════════════════════════════════════

rapport_path = (r"C:\Users\LENOVO\Downloads\project-sara-aham"
                r"\distributed-database-docker-project-main"
                r"\rapport\rapport_BDD_distribuees (1).docx")

doc_r = Document(rapport_path)

# Ajouter entrée TDM après section 21
for para in doc_r.paragraphs:
    if 'Correction de la redondance' in para.text or 'redondance' in para.text.lower():
        new_toc = doc_r.add_paragraph()
        para._p.addnext(new_toc._p)
        r = new_toc.add_run(
            "  22. Correction de la faille de doublons sur Site1 et Site2")
        r.font.size = Pt(11)
        break

doc_r.add_page_break()
build_section(doc_r, 22)
doc_r.save(rapport_path)
print("[OK] Rapport mis a jour : rapport_BDD_distribuees (1).docx")


# ═══════════════════════════════════════════════════════════════════════════════
# 2. Explication
# ═══════════════════════════════════════════════════════════════════════════════

expl_path = (r"C:\Users\LENOVO\Downloads\project-sara-aham"
             r"\distributed-database-docker-project-main"
             r"\rapport\Explication_PPT_BDD_AhlamNasri_SaraHimmad.docx")

doc_e = Document(expl_path)

for para in doc_e.paragraphs:
    if 'redondance' in para.text.lower() or 'TRUNCATE' in para.text:
        new_plan = doc_e.add_paragraph(style='List Bullet')
        para._p.addnext(new_plan._p)
        r = new_plan.add_run(
            "13 · Correction de la faille de doublons sur Site1 et Site2")
        r.font.size = Pt(10.5)
        break

doc_e.add_page_break()
build_section(doc_e, 13)
doc_e.save(expl_path)
print("[OK] Explication mise a jour : Explication_PPT_BDD_AhlamNasri_SaraHimmad.docx")

print("\nMise a jour terminee.")
