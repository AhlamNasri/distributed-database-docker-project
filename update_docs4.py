#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Ajoute la section de synthèse finale des 3 fixes dans les deux documents Word.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLEU_FONCE = RGBColor(0x0D, 0x1B, 0x3E)
ORANGE     = RGBColor(0xE8, 0x6B, 0x1A)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_para(doc, text='', bold=False, size=11, color=None,
             space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p

def add_divider(doc, color='E86B1A'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '8')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def add_cell_text(cell, text, bold=False, size=10, color=None,
                  align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color

def build_synthese(doc, section_number):
    # ── Titre ─────────────────────────────────────────────────────────────────
    h = doc.add_heading(
        f'{section_number}. SYNTHESE DES CORRECTIONS — TABLEAU RECAPITULATIF',
        level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after  = Pt(8)
    for run in h.runs:
        run.font.color.rgb = BLEU_FONCE
    add_divider(doc)

    add_para(doc,
        "Les trois problèmes identifiés lors de l'analyse du système ont été "
        "corrigés de manière complémentaire. Ce tableau récapitule l'ensemble "
        "des corrections apportées, leurs causes racines, les fichiers modifiés "
        "et les résultats de validation obtenus.")

    add_para(doc, '')

    # ── Tableau principal 3 fixes ──────────────────────────────────────────────
    COL_W = [Cm(2.2), Cm(3.8), Cm(3.8), Cm(3.2), Cm(3.2)]
    headers = ["Fix", "Problème", "Solution", "Fichiers modifiés", "Validation"]

    tbl = doc.add_table(rows=4, cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (w, h_txt) in enumerate(zip(COL_W, headers)):
        tbl.columns[i].width = w
        cell = tbl.cell(0, i)
        set_cell_bg(cell, '0D1B3E')
        add_cell_text(cell, h_txt, bold=True, size=10,
                      color=RGBColor(0xFF, 0xFF, 0xFF))

    data = [
        {
            "fix":      "Fix #1\nTrigger\nCentral",
            "prob":     "INSERT direct dans LIGNECOMMANDES sur le Central : "
                        "le trigger AFTER laissait la ligne s'enregistrer avant "
                        "de la router vers le bon site, créant un doublon.",
            "sol":      "Remplacement du trigger AFTER par un trigger BEFORE "
                        "qui lève immédiatement ORA-20300 et annule l'opération. "
                        "Seule la vue V_LIGNECOMMANDES_ROUTAGE est autorisée.",
            "files":    "docker/\nrouting_central.sh",
            "valid":    "ORA-20300 levée sur insert direct ✅\nInsert via vue routé correctement ✅",
            "bg_fix":   "FFF3E0",
            "bg_other": "FFFDE7",
        },
        {
            "fix":      "Fix #2\nTRUNCATE\nInitial",
            "prob":     "Les 25 lignes initiales de LIGNECOMMANDES étaient "
                        "présentes sur les 3 noeuds simultanément après l'init. "
                        "Les fragments étaient créés par SELECT mais la table "
                        "source n'était jamais vidée.",
            "sol":      "Ajout de TRUNCATE TABLE LIGNECOMMANDES à la fin de "
                        "chaque script de fragmentation, après la création des "
                        "tables LigneCommandes1 et LigneCommandes2.",
            "files":    "site1/scenario2/\nfragments/\nsite1_fragments.sh\n\n"
                        "site2/scenario2/\nfragments/\nsite2_fragments.sh\n\n"
                        "docker/\nrouting_central.sh",
            "valid":    "Central = 0 ligne ✅\nSite1 fragment = 12 lignes ✅\n"
                        "Site2 fragment = 13 lignes ✅\nVue globale = 25 total ✅",
            "bg_fix":   "E8F5E9",
            "bg_other": "F1F8E9",
        },
        {
            "fix":      "Fix #3\nTriggers\nSite1/Site2",
            "prob":     "Fix #1 ne couvrait que le Central. Sur Site1 et Site2, "
                        "LIGNECOMMANDES était vide mais un INSERT direct "
                        "s'exécutait sans erreur, stockant la donnée hors "
                        "fragment et hors vue globale.",
            "sol":      "Ajout d'un trigger BEFORE INSERT OR UPDATE OR DELETE "
                        "sur LIGNECOMMANDES dans les fichiers de triggers de "
                        "Site1 (ORA-20301) et Site2 (ORA-20302).",
            "files":    "site1/scenario2/\ntriggers/\nsite1_triggers.sh\n\n"
                        "site2/scenario2/\ntriggers/\nsite2_triggers.sh",
            "valid":    "ORA-20301 sur Site1 ✅\nORA-20302 sur Site2 ✅\n"
                        "3 noeuds protégés ✅",
            "bg_fix":   "E3F2FD",
            "bg_other": "EDE7F6",
        },
    ]

    for ri, row in enumerate(data):
        cells = [tbl.cell(ri + 1, ci) for ci in range(5)]
        set_cell_bg(cells[0], row["bg_fix"])
        add_cell_text(cells[0], row["fix"], bold=True, size=9.5,
                      color=BLEU_FONCE)
        for ci, key in enumerate(["prob", "sol", "files", "valid"]):
            set_cell_bg(cells[ci + 1], row["bg_other"])
            add_cell_text(cells[ci + 1], row[key], bold=False, size=9,
                          align=WD_ALIGN_PARAGRAPH.LEFT)

    add_para(doc, '')

    # ── Tableau de protection par noeud ────────────────────────────────────────
    add_para(doc, 'Protection par noeud après les 3 fixes :',
             bold=True, size=11, color=ORANGE, space_after=4)

    tbl2 = doc.add_table(rows=5, cols=4)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h_txt in enumerate(["Noeud", "Table protégée", "Mécanisme", "Code erreur"]):
        cell = tbl2.cell(0, i)
        set_cell_bg(cell, '0D4E2E')
        add_cell_text(cell, h_txt, bold=True, size=10,
                      color=RGBColor(0xFF, 0xFF, 0xFF))

    prot_rows = [
        ["Central", "LIGNECOMMANDES",           "BEFORE trigger (Fix #1)",          "ORA-20300",  "FFF8E1"],
        ["Site1",   "LIGNECOMMANDES",           "BEFORE trigger (Fix #3)",          "ORA-20301",  "E8F5E9"],
        ["Site2",   "LIGNECOMMANDES",           "BEFORE trigger (Fix #3)",          "ORA-20302",  "E8F5E9"],
        ["Central", "V_LIGNECOMMANDES_ROUTAGE", "INSTEAD OF trigger (routage auto)","—",          "E3F2FD"],
    ]

    for ri, row in enumerate(prot_rows):
        for ci, val in enumerate(row[:4]):
            cell = tbl2.cell(ri + 1, ci)
            set_cell_bg(cell, row[4])
            add_cell_text(cell, val, bold=(ci == 0), size=9.5)

    add_para(doc, '')

    # ── Flux final ────────────────────────────────────────────────────────────
    add_para(doc, 'Flux d\'insertion garanti après les 3 fixes :',
             bold=True, size=11, color=ORANGE, space_after=4)

    add_para(doc,
        "Toute tentative d'écriture directe dans LIGNECOMMANDES sur l'un "
        "des 3 noeuds est immédiatement rejetée. L'unique chemin valide est :",
        size=10)

    flux_tbl = doc.add_table(rows=1, cols=5)
    flux_tbl.style = 'Table Grid'
    flux_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    etapes = [
        ("Application\nclient", "E3F2FD"),
        ("V_LIGNECOMMANDES\n_ROUTAGE\n(Central)", "FFF8E1"),
        ("INSTEAD OF\ntrigger", "FFF3E0"),
        ("LigneCommandes1\n(QUANTITE >= 100)", "E8F5E9"),
        ("LigneCommandes2\n(QUANTITE < 100)", "FCE4EC"),
    ]

    for ci, (txt, bg) in enumerate(etapes):
        cell = flux_tbl.cell(0, ci)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(txt)
        r.bold = (ci in [1, 3, 4])
        r.font.size = Pt(8.5)

    add_para(doc, '')
    add_para(doc,
        "Ce flux est identique à celui décrit dans la conception initiale du système. "
        "Les trois corrections n'ajoutent pas de nouvelle logique — elles ferment les "
        "chemins d'accès qui contournaient involontairement ce flux.",
        size=10)

    add_para(doc, '')


# ═══════════════════════════════════════════════════════════════════════════════
# 1. Rapport
# ═══════════════════════════════════════════════════════════════════════════════

rapport_path = (r"C:\Users\LENOVO\Downloads\project-sara-aham"
                r"\distributed-database-docker-project-main"
                r"\rapport\rapport_BDD_distribuees (1).docx")

doc_r = Document(rapport_path)

# Ajouter entrée TDM après section 22
for para in doc_r.paragraphs:
    if 'faille de doublons' in para.text.lower() or 'Site1 et Site2' in para.text:
        new_toc = doc_r.add_paragraph()
        para._p.addnext(new_toc._p)
        r = new_toc.add_run(
            "  23. Synthese des corrections — Tableau récapitulatif")
        r.bold = True
        r.font.size = Pt(11)
        break

doc_r.add_page_break()
build_synthese(doc_r, 23)
doc_r.save(rapport_path)
print("[OK] Rapport mis a jour.")


# ═══════════════════════════════════════════════════════════════════════════════
# 2. Explication
# ═══════════════════════════════════════════════════════════════════════════════

expl_path = (r"C:\Users\LENOVO\Downloads\project-sara-aham"
             r"\distributed-database-docker-project-main"
             r"\rapport\Explication_PPT_BDD_AhlamNasri_SaraHimmad.docx")

doc_e = Document(expl_path)

for para in doc_e.paragraphs:
    if 'faille de doublons' in para.text.lower() or 'Site1 et Site2' in para.text:
        new_plan = doc_e.add_paragraph(style='List Bullet')
        para._p.addnext(new_plan._p)
        r = new_plan.add_run(
            "14 · Synthese des corrections — Tableau récapitulatif")
        r.bold = True
        r.font.size = Pt(10.5)
        break

doc_e.add_page_break()
build_synthese(doc_e, 14)
doc_e.save(expl_path)
print("[OK] Explication mise a jour.")

print("\nTermine.")
